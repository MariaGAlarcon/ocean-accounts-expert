#!/usr/bin/env python3
"""Generate GOAP-styled Excel workbooks and Word documents for the clinic."""

import math
from pathlib import Path

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---- GOAP Styles ----
GREEN = "3B9C7B"
TEAL = "0A5455"
GRAY = "404040"
TEXT_COLOR = "30302F"
WHITE = "FFFFFF"
YELLOW = "FFF2CC"
LIGHT_GRAY = "F2F2F2"
MINT = "D4EEE5"

green_fill = PatternFill(start_color=GREEN, end_color=GREEN, fill_type="solid")
teal_fill = PatternFill(start_color=TEAL, end_color=TEAL, fill_type="solid")
yellow_fill = PatternFill(start_color=YELLOW, end_color=YELLOW, fill_type="solid")
mint_fill = PatternFill(start_color=MINT, end_color=MINT, fill_type="solid")
gray_fill = PatternFill(start_color=LIGHT_GRAY, end_color=LIGHT_GRAY, fill_type="solid")

bdr = Border(
    left=Side("thin", GRAY), right=Side("thin", GRAY),
    top=Side("thin", GRAY), bottom=Side("thin", GRAY),
)
hdr_font = Font(name="Arial", size=10, bold=True, color=WHITE)
data_font = Font(name="Arial", size=10, color=TEXT_COLOR)
label_font = Font(name="Arial", size=10, bold=True, color=TEAL)
title_font = Font(name="Arial", size=14, bold=True, color=TEAL)
subtitle_font = Font(name="Arial", size=11, bold=False, color=GREEN)
input_font = Font(name="Arial", size=10, color=TEXT_COLOR)
answer_font = Font(name="Arial", size=10, bold=True, color="006100")
answer_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")


def apply_header(cell):
    cell.font = hdr_font
    cell.fill = green_fill
    cell.border = bdr
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def apply_row_header(cell):
    cell.font = Font(name="Arial", size=10, bold=True, color=WHITE)
    cell.fill = teal_fill
    cell.border = bdr
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)


def apply_data(cell, align="left"):
    cell.font = data_font
    cell.border = bdr
    cell.alignment = Alignment(horizontal=align, vertical="center")


def apply_input(cell):
    cell.fill = yellow_fill
    cell.font = input_font
    cell.border = bdr
    cell.alignment = Alignment(horizontal="right", vertical="center")


def apply_answer(cell):
    cell.font = answer_font
    cell.fill = answer_fill
    cell.border = bdr
    cell.alignment = Alignment(horizontal="right", vertical="center")


def set_col_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w


# =========================================================================
# 1. EXTENT ACCOUNT TEMPLATE
# =========================================================================
def generate_extent_template(outdir):
    wb = openpyxl.Workbook()

    # -- Sheet 1: Extent Account --
    ws = wb.active
    ws.title = "Extent Account"
    ws.sheet_properties.tabColor = GREEN

    # Title
    ws["A1"] = "SEEA EA Ecosystem Extent Account"
    ws["A1"].font = title_font
    ws["A2"] = "Accounting area:"
    ws["A2"].font = label_font
    ws["B2"].fill = yellow_fill
    ws["B2"].border = bdr
    ws["A3"] = "Opening year:"
    ws["A3"].font = label_font
    ws["B3"].fill = yellow_fill
    ws["B3"].border = bdr
    ws["A4"] = "Closing year:"
    ws["A4"].font = label_font
    ws["B4"].fill = yellow_fill
    ws["B4"].border = bdr

    # Headers row 6
    headers = ["", "Coral reefs\n(M1.3)", "Seagrass\n(M1.1)", "Mangroves\n(MFT1.2)",
               "Ecosystem 4\n(edit)", "Other", "Total"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=6, column=c, value=h)
        apply_header(cell)

    rows = [
        "Opening extent (ha)",
        "Additions: managed expansion",
        "Additions: natural expansion",
        "Additions: reclassification",
        "Total additions",
        "Reductions: managed reduction",
        "Reductions: natural reduction",
        "Reductions: reclassification",
        "Total reductions",
        "Net change",
        "Closing extent (ha)",
    ]
    for r, label in enumerate(rows, 7):
        cell = ws.cell(row=r, column=1, value=label)
        apply_row_header(cell)
        for c in range(2, 8):
            cell = ws.cell(row=r, column=c)
            apply_input(cell)

    # Formulas for Total additions (row 11), Total reductions (row 15), Net change (row 16), Closing (row 17)
    for c in range(2, 8):
        col = get_column_letter(c)
        ws.cell(row=11, column=c).value = f"={col}8+{col}9+{col}10"
        ws.cell(row=15, column=c).value = f"={col}12+{col}13+{col}14"
        ws.cell(row=16, column=c).value = f"={col}11-{col}15"
        ws.cell(row=17, column=c).value = f"={col}7+{col}16"
    # Total column (G) sums B-F
    for r in range(7, 18):
        col_g = get_column_letter(7)
        ws.cell(row=r, column=7).value = f"=SUM(B{r}:F{r})"

    set_col_widths(ws, [32, 16, 16, 16, 16, 16, 16])

    # -- Sheet 2: Change Matrix --
    ws2 = wb.create_sheet("Change Matrix")
    ws2.sheet_properties.tabColor = TEAL

    ws2["A1"] = "Ecosystem Type Change Matrix (hectares)"
    ws2["A1"].font = title_font

    labels = ["Coral reefs", "Seagrass", "Mangroves", "Ecosystem 4", "Other"]
    # Header row
    ws2.cell(row=3, column=1, value="Opening \\ Closing").font = hdr_font
    ws2.cell(row=3, column=1).fill = teal_fill
    ws2.cell(row=3, column=1).border = bdr
    for c, lb in enumerate(labels, 2):
        cell = ws2.cell(row=3, column=c, value=lb)
        apply_header(cell)
    cell = ws2.cell(row=3, column=7, value="Opening total")
    apply_header(cell)

    for r, lb in enumerate(labels, 4):
        cell = ws2.cell(row=r, column=1, value=lb)
        apply_row_header(cell)
        for c in range(2, 7):
            cell = ws2.cell(row=r, column=c)
            apply_input(cell)
        # Row total
        ws2.cell(row=r, column=7).value = f"=SUM(B{r}:F{r})"
        apply_data(ws2.cell(row=r, column=7), "right")

    # Closing total row
    r_total = 4 + len(labels)
    ws2.cell(row=r_total, column=1, value="Closing total").font = Font(
        name="Arial", size=10, bold=True, color=WHITE)
    ws2.cell(row=r_total, column=1).fill = teal_fill
    ws2.cell(row=r_total, column=1).border = bdr
    for c in range(2, 7):
        ws2.cell(row=r_total, column=c).value = f"=SUM({get_column_letter(c)}4:{get_column_letter(c)}8)"
        apply_data(ws2.cell(row=r_total, column=c), "right")

    set_col_widths(ws2, [20, 16, 16, 16, 16, 16, 16])

    # -- Sheet 3: Instructions --
    ws3 = wb.create_sheet("Instructions")
    ws3["A1"] = "How to use this template"
    ws3["A1"].font = title_font
    instructions = [
        "1. Fill in the yellow cells with your data.",
        "2. On the Extent Account sheet, enter area values in hectares.",
        "3. Total additions, reductions, net change, and closing extent calculate automatically.",
        "4. On the Change Matrix sheet, enter the cross-tabulation of opening vs closing classifications.",
        "5. Row totals (opening extent per type) and column totals (closing extent per type) calculate automatically.",
        "6. The Total column on the Extent Account sums across all ecosystem types.",
        "7. Adapt the ecosystem type names and IUCN GET codes to your accounting area.",
        "8. Delete Ecosystem 4 column if not needed, or add more columns as required.",
    ]
    for i, txt in enumerate(instructions, 3):
        ws3.cell(row=i, column=1, value=txt).font = data_font

    ws3.column_dimensions["A"].width = 80

    path = outdir / "extent" / "extent_account_template.xlsx"
    wb.save(path)
    print(f"  Created: {path}")


# =========================================================================
# 2. CONDITION ACCOUNT TEMPLATE + EXERCISE
# =========================================================================
def generate_condition_template(outdir):
    wb = openpyxl.Workbook()

    # -- Sheet 1: Condition Account --
    ws = wb.active
    ws.title = "Condition Account"
    ws.sheet_properties.tabColor = GREEN

    ws["A1"] = "SEEA EA Ecosystem Condition Account"
    ws["A1"].font = title_font
    ws["A2"] = "Ecosystem type:"
    ws["A2"].font = label_font
    ws["B2"].fill = yellow_fill
    ws["B2"].border = bdr
    ws["A3"] = "Accounting area:"
    ws["A3"].font = label_font
    ws["B3"].fill = yellow_fill
    ws["B3"].border = bdr
    ws["A4"] = "Opening year:"
    ws["A4"].font = label_font
    ws["B4"].fill = yellow_fill
    ws["B4"].border = bdr
    ws["C4"] = "Closing year:"
    ws["C4"].font = label_font
    ws["D4"].fill = yellow_fill
    ws["D4"].border = bdr

    headers = ["Indicator", "Unit", "Direction", "Reference\nLevel",
               "Opening\nValue", "Opening\nCI", "Closing\nValue", "Closing\nCI",
               "Change\nin CI", "Interpretation"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=6, column=c, value=h)
        apply_header(cell)

    # Pre-fill 8 indicator rows
    for r in range(7, 15):
        ws.cell(row=r, column=1).border = bdr
        ws.cell(row=r, column=1).font = data_font
        for c in range(2, 11):
            cell = ws.cell(row=r, column=c)
            if c <= 4:
                apply_data(cell)
            else:
                apply_input(cell)
        # CI formulas
        # Opening CI (col 6): if direction="Higher is better" => E/D, else => 1-(E/D)
        # We'll use a helper: user enters direction in col 3
        col_e = f"E{r}"
        col_d = f"D{r}"
        ws.cell(row=r, column=6).value = (
            f'=IF({col_e}="","",IF(C{r}="Higher is better",'
            f'MIN({col_e}/{col_d},1),MAX(1-{col_e}/{col_d},0)))'
        )
        col_g = f"G{r}"
        ws.cell(row=r, column=8).value = (
            f'=IF({col_g}="","",IF(C{r}="Higher is better",'
            f'MIN({col_g}/{col_d},1),MAX(1-{col_g}/{col_d},0)))'
        )
        # Change in CI
        ws.cell(row=r, column=9).value = f'=IF(OR(F{r}="",H{r}=""),"",H{r}-F{r})'
        # Interpretation
        ws.cell(row=r, column=10).value = (
            f'=IF(I{r}="","",IF(I{r}>0.01,"Improving",IF(I{r}<-0.01,"Declining","Stable")))'
        )

    set_col_widths(ws, [28, 10, 18, 12, 12, 12, 12, 12, 12, 16])

    # -- Sheet 2: Raw Data --
    ws2 = wb.create_sheet("Raw Data")
    ws2.sheet_properties.tabColor = TEAL

    ws2["A1"] = "Site-Level Indicator Data"
    ws2["A1"].font = title_font
    ws2["A2"] = "Paste or enter your site-level measurements below."
    ws2["A2"].font = subtitle_font

    raw_headers = ["Year", "Site ID", "Geographic Unit", "Indicator 1", "Indicator 2",
                   "Indicator 3", "Indicator 4", "Indicator 5", "Indicator 6"]
    for c, h in enumerate(raw_headers, 1):
        cell = ws2.cell(row=4, column=c, value=h)
        apply_header(cell)

    for r in range(5, 55):
        for c in range(1, 10):
            cell = ws2.cell(row=r, column=c)
            apply_data(cell, "right" if c >= 4 else "left")

    set_col_widths(ws2, [10, 14, 20, 14, 14, 14, 14, 14, 14])

    # -- Sheet 3: Reference Levels --
    ws3 = wb.create_sheet("Reference Levels")
    ref_headers = ["Indicator", "Unit", "Direction", "Reference Level", "Source"]
    for c, h in enumerate(ref_headers, 1):
        cell = ws3.cell(row=1, column=c, value=h)
        apply_header(cell)
    for r in range(2, 10):
        for c in range(1, 6):
            ws3.cell(row=r, column=c).border = bdr
            ws3.cell(row=r, column=c).font = data_font

    set_col_widths(ws3, [28, 10, 18, 14, 40])

    path = outdir / "condition" / "condition_account_template.xlsx"
    wb.save(path)
    print(f"  Created: {path}")


def generate_condition_exercise(outdir):
    wb = openpyxl.Workbook()

    ws = wb.active
    ws.title = "Exercise"
    ws.sheet_properties.tabColor = GREEN

    ws["A1"] = "Condition Account Exercise"
    ws["A1"].font = title_font
    ws["A2"] = "Coral reef condition -- 4 sites, 3 indicators, 2 periods"
    ws["A2"].font = subtitle_font

    # Reference levels
    ws["A4"] = "Reference Levels"
    ws["A4"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    ref_h = ["Indicator", "Reference", "Direction"]
    for c, h in enumerate(ref_h, 1):
        cell = ws.cell(row=5, column=c, value=h)
        apply_header(cell)
    refs = [
        ("Live coral cover (%)", 50, "Higher is better"),
        ("Fleshy macroalgae (%)", 60, "Higher is worse"),
        ("Fish biomass (kg/ha)", 500, "Higher is better"),
    ]
    for r, (name, ref, d) in enumerate(refs, 6):
        ws.cell(row=r, column=1, value=name).font = data_font
        ws.cell(row=r, column=1).border = bdr
        ws.cell(row=r, column=2, value=ref).font = data_font
        ws.cell(row=r, column=2).border = bdr
        ws.cell(row=r, column=2).alignment = Alignment(horizontal="right")
        ws.cell(row=r, column=3, value=d).font = data_font
        ws.cell(row=r, column=3).border = bdr

    # Opening data
    ws["A10"] = "Opening Data (2019)"
    ws["A10"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    data_h = ["Site", "Coral cover (%)", "Macroalgae (%)", "Fish biomass (kg/ha)"]
    for c, h in enumerate(data_h, 1):
        cell = ws.cell(row=11, column=c, value=h)
        apply_header(cell)

    opening = [
        ("Site A", 28.5, 22.0, 180),
        ("Site B", 35.2, 18.5, 245),
        ("Site C", 18.7, 38.0, 120),
        ("Site D", 42.1, 12.0, 310),
    ]
    for r, (site, cc, ma, fb) in enumerate(opening, 12):
        ws.cell(row=r, column=1, value=site)
        apply_row_header(ws.cell(row=r, column=1))
        for c, v in enumerate([cc, ma, fb], 2):
            cell = ws.cell(row=r, column=c, value=v)
            apply_data(cell, "right")

    # Closing data
    ws["A17"] = "Closing Data (2024)"
    ws["A17"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    for c, h in enumerate(data_h, 1):
        cell = ws.cell(row=18, column=c, value=h)
        apply_header(cell)

    closing = [
        ("Site A", 22.0, 28.5, 160),
        ("Site B", 30.8, 24.0, 220),
        ("Site C", 15.2, 42.0, 105),
        ("Site D", 38.5, 15.0, 290),
    ]
    for r, (site, cc, ma, fb) in enumerate(closing, 19):
        ws.cell(row=r, column=1, value=site)
        apply_row_header(ws.cell(row=r, column=1))
        for c, v in enumerate([cc, ma, fb], 2):
            cell = ws.cell(row=r, column=c, value=v)
            apply_data(cell, "right")

    # Part 1: Normalize -- yellow input cells
    ws["A24"] = "Part 1: Calculate Condition Index (fill yellow cells)"
    ws["A24"].font = Font(name="Arial", size=12, bold=True, color=TEAL)

    ci_h = ["Site", "Coral CI", "Macroalgae CI", "Fish CI"]
    ws["A25"] = "Opening (2019)"
    ws["A25"].font = label_font
    for c, h in enumerate(ci_h, 1):
        cell = ws.cell(row=26, column=c, value=h)
        apply_header(cell)
    for r in range(27, 31):
        site_r = r - 27 + 12  # row in opening data
        ws.cell(row=r, column=1, value=ws.cell(row=site_r, column=1).value)
        apply_row_header(ws.cell(row=r, column=1))
        for c in range(2, 5):
            apply_input(ws.cell(row=r, column=c))

    ws["A32"] = "Closing (2024)"
    ws["A32"].font = label_font
    for c, h in enumerate(ci_h, 1):
        cell = ws.cell(row=33, column=c, value=h)
        apply_header(cell)
    for r in range(34, 38):
        site_r = r - 34 + 19
        ws.cell(row=r, column=1, value=ws.cell(row=site_r, column=1).value)
        apply_row_header(ws.cell(row=r, column=1))
        for c in range(2, 5):
            apply_input(ws.cell(row=r, column=c))

    # Part 2: Aggregate
    ws["A39"] = "Part 2: Calculate Mean CI (fill yellow cells)"
    ws["A39"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    agg_h = ["", "Coral CI", "Macroalgae CI", "Fish CI"]
    for c, h in enumerate(agg_h, 1):
        cell = ws.cell(row=40, column=c, value=h)
        apply_header(cell)
    ws.cell(row=41, column=1, value="Opening Mean CI")
    apply_row_header(ws.cell(row=41, column=1))
    ws.cell(row=42, column=1, value="Closing Mean CI")
    apply_row_header(ws.cell(row=42, column=1))
    ws.cell(row=43, column=1, value="Change in CI")
    apply_row_header(ws.cell(row=43, column=1))
    ws.cell(row=44, column=1, value="Direction")
    apply_row_header(ws.cell(row=44, column=1))
    for r in range(41, 45):
        for c in range(2, 5):
            apply_input(ws.cell(row=r, column=c))

    set_col_widths(ws, [28, 16, 16, 20])

    # -- Answer Sheet --
    ws_a = wb.create_sheet("Answers")
    ws_a.sheet_properties.tabColor = "006100"
    ws_a["A1"] = "Answer Key"
    ws_a["A1"].font = title_font

    ws_a["A3"] = "Opening CI"
    ws_a["A3"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    for c, h in enumerate(ci_h, 1):
        cell = ws_a.cell(row=4, column=c, value=h)
        apply_header(cell)

    answers_open = [
        ("Site A", 0.57, 0.63, 0.36),
        ("Site B", 0.70, 0.69, 0.49),
        ("Site C", 0.37, 0.37, 0.24),
        ("Site D", 0.84, 0.80, 0.62),
    ]
    for r, (site, cc, ma, fb) in enumerate(answers_open, 5):
        ws_a.cell(row=r, column=1, value=site)
        apply_row_header(ws_a.cell(row=r, column=1))
        for c, v in enumerate([cc, ma, fb], 2):
            cell = ws_a.cell(row=r, column=c, value=v)
            apply_answer(cell)
            cell.number_format = "0.00"

    ws_a["A10"] = "Closing CI"
    ws_a["A10"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    for c, h in enumerate(ci_h, 1):
        cell = ws_a.cell(row=11, column=c, value=h)
        apply_header(cell)

    answers_close = [
        ("Site A", 0.44, 0.53, 0.32),
        ("Site B", 0.62, 0.60, 0.44),
        ("Site C", 0.30, 0.30, 0.21),
        ("Site D", 0.77, 0.75, 0.58),
    ]
    for r, (site, cc, ma, fb) in enumerate(answers_close, 12):
        ws_a.cell(row=r, column=1, value=site)
        apply_row_header(ws_a.cell(row=r, column=1))
        for c, v in enumerate([cc, ma, fb], 2):
            cell = ws_a.cell(row=r, column=c, value=v)
            apply_answer(cell)
            cell.number_format = "0.00"

    ws_a["A17"] = "Condition Account Summary"
    ws_a["A17"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    sum_h = ["Indicator", "Opening CI", "Closing CI", "Change", "Direction"]
    for c, h in enumerate(sum_h, 1):
        cell = ws_a.cell(row=18, column=c, value=h)
        apply_header(cell)
    summary = [
        ("Coral cover", 0.62, 0.53, -0.09, "Declining"),
        ("Macroalgae", 0.62, 0.55, -0.07, "Declining"),
        ("Fish biomass", 0.43, 0.39, -0.04, "Declining"),
    ]
    for r, (ind, o, cl, ch, d) in enumerate(summary, 19):
        ws_a.cell(row=r, column=1, value=ind)
        apply_row_header(ws_a.cell(row=r, column=1))
        for c, v in enumerate([o, cl, ch], 2):
            cell = ws_a.cell(row=r, column=c, value=v)
            apply_answer(cell)
            cell.number_format = "0.00"
        ws_a.cell(row=r, column=5, value=d)
        apply_data(ws_a.cell(row=r, column=5))

    set_col_widths(ws_a, [28, 16, 16, 16, 16])

    path = outdir / "condition" / "condition_exercise.xlsx"
    wb.save(path)
    print(f"  Created: {path}")


# =========================================================================
# 3. SERVICES ACCOUNT TEMPLATE + EXERCISE
# =========================================================================
def generate_services_template(outdir):
    wb = openpyxl.Workbook()

    ws = wb.active
    ws.title = "Physical Supply"
    ws.sheet_properties.tabColor = GREEN

    ws["A1"] = "SEEA EA Ecosystem Service Supply Account"
    ws["A1"].font = title_font
    ws["A2"] = "Accounting area:"
    ws["A2"].font = label_font
    ws["B2"].fill = yellow_fill
    ws["B2"].border = bdr
    ws["A3"] = "Accounting year:"
    ws["A3"].font = label_font
    ws["B3"].fill = yellow_fill
    ws["B3"].border = bdr

    headers = ["Service", "Unit", "Coral reefs", "Seagrass", "Mangroves", "Total"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=5, column=c, value=h)
        apply_header(cell)

    services = [
        ("Fish provisioning", "kg/yr"),
        ("Carbon sequestration", "Mg CO2/yr"),
        ("Coastal protection", "m coastline"),
        ("Nursery habitat", "kg biomass/yr"),
        ("Recreation", "visitors/yr"),
        ("Gleaning", "hours/yr"),
    ]
    for r, (svc, unit) in enumerate(services, 6):
        ws.cell(row=r, column=1, value=svc)
        apply_row_header(ws.cell(row=r, column=1))
        ws.cell(row=r, column=2, value=unit)
        apply_data(ws.cell(row=r, column=2))
        for c in range(3, 6):
            apply_input(ws.cell(row=r, column=c))
        ws.cell(row=r, column=6).value = f"=SUM(C{r}:E{r})"
        apply_data(ws.cell(row=r, column=6), "right")

    set_col_widths(ws, [24, 16, 16, 16, 16, 16])

    # -- Sheet 2: Monetary Supply --
    ws2 = wb.create_sheet("Monetary Supply")
    ws2.sheet_properties.tabColor = TEAL

    ws2["A1"] = "Monetary Ecosystem Service Supply (USD/yr)"
    ws2["A1"].font = title_font

    m_headers = ["Service", "Valuation method", "Coral reefs", "Seagrass", "Mangroves", "Total"]
    for c, h in enumerate(m_headers, 1):
        cell = ws2.cell(row=3, column=c, value=h)
        apply_header(cell)

    m_services = [
        ("Fish provisioning", "Resource rent"),
        ("Carbon sequestration", "Social cost of carbon"),
        ("Coastal protection", "Replacement cost"),
        ("Nursery habitat", "Productivity change"),
        ("Recreation", "Direct expenditure"),
        ("Gleaning", "Equivalent wage"),
        ("TOTAL", ""),
    ]
    for r, (svc, method) in enumerate(m_services, 4):
        ws2.cell(row=r, column=1, value=svc)
        if svc == "TOTAL":
            apply_row_header(ws2.cell(row=r, column=1))
            ws2.cell(row=r, column=2, value="")
            apply_row_header(ws2.cell(row=r, column=2))
            for c in range(3, 7):
                ws2.cell(row=r, column=c).value = f"=SUM({get_column_letter(c)}4:{get_column_letter(c)}9)"
                apply_data(ws2.cell(row=r, column=c), "right")
                ws2.cell(row=r, column=c).font = Font(name="Arial", size=10, bold=True, color=TEXT_COLOR)
        else:
            apply_row_header(ws2.cell(row=r, column=1))
            ws2.cell(row=r, column=2, value=method)
            apply_data(ws2.cell(row=r, column=2))
            for c in range(3, 6):
                apply_input(ws2.cell(row=r, column=c))
            ws2.cell(row=r, column=6).value = f"=SUM(C{r}:E{r})"
            apply_data(ws2.cell(row=r, column=6), "right")

    set_col_widths(ws2, [24, 22, 16, 16, 16, 16])

    # -- Sheet 3: Calculation helpers --
    ws3 = wb.create_sheet("Calculations")
    ws3["A1"] = "Quick Calculation Helpers"
    ws3["A1"].font = title_font

    # Carbon calculator
    ws3["A3"] = "Carbon Sequestration Calculator"
    ws3["A3"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    calc_h = ["Ecosystem", "Extent (ha)", "NCP rate\n(Mg CO2/ha/yr)", "Physical supply\n(Mg CO2/yr)",
              "Carbon price\n(USD/Mg CO2)", "Value\n(USD/yr)"]
    for c, h in enumerate(calc_h, 1):
        cell = ws3.cell(row=4, column=c, value=h)
        apply_header(cell)

    eco_defaults = [("Mangroves", 3.7), ("Seagrass", 1.3)]
    for r, (eco, ncp) in enumerate(eco_defaults, 5):
        ws3.cell(row=r, column=1, value=eco)
        apply_row_header(ws3.cell(row=r, column=1))
        apply_input(ws3.cell(row=r, column=2))  # extent
        ws3.cell(row=r, column=3, value=ncp)
        apply_data(ws3.cell(row=r, column=3), "right")
        ws3.cell(row=r, column=4).value = f"=B{r}*C{r}"
        apply_data(ws3.cell(row=r, column=4), "right")
        ws3.cell(row=r, column=5, value=51)
        apply_data(ws3.cell(row=r, column=5), "right")
        ws3.cell(row=r, column=6).value = f"=D{r}*E{r}"
        apply_data(ws3.cell(row=r, column=6), "right")

    # Resource rent calculator
    ws3["A9"] = "Resource Rent Calculator (Fish Provisioning)"
    ws3["A9"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    rr_labels = ["Total catch (kg/yr)", "Average price (USD/kg)", "Gross revenue (USD/yr)",
                 "Labour costs (USD/yr)", "Fuel costs (USD/yr)", "Gear and maintenance (USD/yr)",
                 "Capital depreciation (USD/yr)", "Total costs (USD/yr)", "Resource rent (USD/yr)"]
    for r, lb in enumerate(rr_labels, 10):
        ws3.cell(row=r, column=1, value=lb)
        ws3.cell(row=r, column=1).font = label_font
        ws3.cell(row=r, column=1).border = bdr
        apply_input(ws3.cell(row=r, column=2))

    ws3.cell(row=12, column=2).value = "=B10*B11"  # Gross revenue
    ws3.cell(row=17, column=2).value = "=SUM(B13:B16)"  # Total costs
    ws3.cell(row=18, column=2).value = "=MAX(B12-B17,0)"  # Resource rent

    set_col_widths(ws3, [32, 16, 18, 18, 18, 16])

    path = outdir / "services" / "services_account_template.xlsx"
    wb.save(path)
    print(f"  Created: {path}")


def generate_services_exercise(outdir):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Exercise"
    ws.sheet_properties.tabColor = GREEN

    ws["A1"] = "Ecosystem Services Exercise"
    ws["A1"].font = title_font
    ws["A2"] = "4 services: fish provisioning, carbon, recreation, coastal protection"
    ws["A2"].font = subtitle_font

    # Given data
    ws["A4"] = "Given Data"
    ws["A4"].font = Font(name="Arial", size=12, bold=True, color=TEAL)
    given = [
        ("Coral reef extent", "1,200 ha"),
        ("Mangrove extent", "350 ha"),
        ("Seagrass extent", "800 ha"),
        ("Reef fish catch", "180,000 kg/yr"),
        ("Average fish price", "USD 3.50/kg"),
        ("Total fishing costs", "USD 420,000/yr"),
        ("Mangrove NCP rate", "3.7 Mg CO2/ha/yr"),
        ("Seagrass NCP rate", "1.3 Mg CO2/ha/yr"),
        ("Social cost of carbon", "USD 51/Mg CO2"),
        ("Reef visitors", "15,000/yr"),
        ("Reef spending per visitor", "USD 85"),
        ("Protected coastline", "12,000 m"),
        ("Seawall replacement cost", "USD 5,000/m"),
        ("Seawall lifespan", "50 years"),
    ]
    for r, (lb, val) in enumerate(given, 5):
        ws.cell(row=r, column=1, value=lb).font = data_font
        ws.cell(row=r, column=1).border = bdr
        ws.cell(row=r, column=2, value=val).font = data_font
        ws.cell(row=r, column=2).border = bdr

    # Calculation sections
    row = 21
    ws.cell(row=row, column=1, value="Part 1: Fish Provisioning").font = Font(
        name="Arial", size=12, bold=True, color=TEAL)
    row += 1
    for lb in ["Gross revenue", "Total costs", "Resource rent"]:
        ws.cell(row=row, column=1, value=lb).font = label_font
        ws.cell(row=row, column=1).border = bdr
        apply_input(ws.cell(row=row, column=2))
        row += 1

    row += 1
    ws.cell(row=row, column=1, value="Part 2: Carbon Sequestration").font = Font(
        name="Arial", size=12, bold=True, color=TEAL)
    row += 1
    for lb in ["Mangrove physical (Mg CO2/yr)", "Mangrove value (USD/yr)",
               "Seagrass physical (Mg CO2/yr)", "Seagrass value (USD/yr)",
               "Total carbon value (USD/yr)"]:
        ws.cell(row=row, column=1, value=lb).font = label_font
        ws.cell(row=row, column=1).border = bdr
        apply_input(ws.cell(row=row, column=2))
        row += 1

    row += 1
    ws.cell(row=row, column=1, value="Part 3: Recreation").font = Font(
        name="Arial", size=12, bold=True, color=TEAL)
    row += 1
    ws.cell(row=row, column=1, value="Recreation value (USD/yr)").font = label_font
    ws.cell(row=row, column=1).border = bdr
    apply_input(ws.cell(row=row, column=2))
    row += 2

    ws.cell(row=row, column=1, value="Part 4: Coastal Protection").font = Font(
        name="Arial", size=12, bold=True, color=TEAL)
    row += 1
    for lb in ["Total replacement cost (USD)", "Annualized value (USD/yr)"]:
        ws.cell(row=row, column=1, value=lb).font = label_font
        ws.cell(row=row, column=1).border = bdr
        apply_input(ws.cell(row=row, column=2))
        row += 1

    set_col_widths(ws, [36, 20])

    # Answer sheet
    ws_a = wb.create_sheet("Answers")
    ws_a.sheet_properties.tabColor = "006100"
    ws_a["A1"] = "Answer Key"
    ws_a["A1"].font = title_font

    answers = [
        ("Fish: Gross revenue", 630000),
        ("Fish: Resource rent", 210000),
        ("Mangrove carbon (Mg CO2/yr)", 1295),
        ("Mangrove carbon value (USD/yr)", 66045),
        ("Seagrass carbon (Mg CO2/yr)", 1040),
        ("Seagrass carbon value (USD/yr)", 53040),
        ("Total carbon value (USD/yr)", 119085),
        ("Recreation value (USD/yr)", 1275000),
        ("Coastal protection total (USD)", 60000000),
        ("Coastal protection annualized (USD/yr)", 1200000),
        ("", ""),
        ("TOTAL ecosystem services (USD/yr)", 2804085),
    ]
    for r, (lb, val) in enumerate(answers, 3):
        ws_a.cell(row=r, column=1, value=lb).font = label_font
        ws_a.cell(row=r, column=1).border = bdr
        cell = ws_a.cell(row=r, column=2, value=val)
        if val:
            apply_answer(cell)
            cell.number_format = "#,##0"

    set_col_widths(ws_a, [40, 20])

    path = outdir / "services" / "services_exercise.xlsx"
    wb.save(path)
    print(f"  Created: {path}")


# =========================================================================
# 4. WORD DOCUMENTS (using GOAP template if available, else plain styled)
# =========================================================================
def goap_table(doc, headers, rows, col_widths=None):
    """Create a GOAP-styled table in a Word document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="3B9C7B"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row_data in enumerate(rows, 1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x30, 0x30, 0x2F)
            # First column teal
            if c == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0A5455"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:left w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:right w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="404040"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    return table


def generate_clinic_guide_docx(outdir):
    """Generate a single GOAP-styled clinic guide Word document."""
    template_path = (
        Path(__file__).parent.parent.parent.parent
        / "goap-templates"
        / "GOAP Word templates"
        / "GOAP-REPORT TEMPLATE.docx"
    )
    if template_path.exists():
        doc = Document(str(template_path))
        # Clear template body
        for p in doc.paragraphs:
            p.clear()
    else:
        doc = Document()

    # Title
    p = doc.add_paragraph()
    run = p.add_run("Ocean Accounting Clinic\nFacilitator Guide")
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0A, 0x54, 0x55)

    p = doc.add_paragraph()
    run = p.add_run("South-South Ocean Accounting Exchange 2026\nBogor, Indonesia")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x3B, 0x9C, 0x7B)

    # Triage section
    h = doc.add_heading("Step 1: Triage", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x54, 0x55)

    p = doc.add_paragraph("Ask the fellow three questions:")
    p.style.font.name = "Arial"
    for q in [
        "1. What account type? (Extent, Condition, Services, or unsure)",
        "2. What data do you have? (None, global datasets, national data, field data)",
        "3. What is your policy question?",
    ]:
        doc.add_paragraph(q, style="List Bullet")

    h = doc.add_heading("Step 2: Direct to pathway", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x54, 0x55)

    goap_table(doc,
        ["They want...", "They have...", "Go to..."],
        [
            ["Extent account", "No data", "Extent Pathway A (global datasets)"],
            ["Extent account", "Satellite imagery or maps", "Extent Pathway B (classification)"],
            ["Condition account", "No field data", "Condition Pathway A (RS proxies)"],
            ["Condition account", "Field survey data", "Condition Pathway B (indicators)"],
            ["Services account", "No data", "Services Pathway A (value transfer)"],
            ["Services account", "Catch data, visitor stats", "Services Pathway B (primary)"],
            ["Don't know", "Any", "Diagnostic tool"],
        ])

    # Account type summaries
    for title, desc in [
        ("Extent Accounts", "Measure ecosystem area (ha) and track changes over time. Foundation for all other accounts."),
        ("Condition Accounts", "Measure ecosystem health relative to reference condition. Condition Index from 0 (degraded) to 1 (reference)."),
        ("Ecosystem Services Accounts", "Quantify benefits to people in physical and monetary terms."),
    ]:
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x0A, 0x54, 0x55)
        p = doc.add_paragraph(desc)
        p.style.font.name = "Arial"

    path = outdir / "clinic_facilitator_guide.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =========================================================================
# MAIN
# =========================================================================
def main():
    outdir = Path(__file__).parent
    print("Generating clinic files...")

    generate_extent_template(outdir)
    generate_condition_template(outdir)
    generate_condition_exercise(outdir)
    generate_services_template(outdir)
    generate_services_exercise(outdir)
    generate_clinic_guide_docx(outdir)

    print("\nDone! All files generated.")


if __name__ == "__main__":
    main()
