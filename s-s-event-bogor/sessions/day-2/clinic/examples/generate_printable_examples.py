#!/usr/bin/env python3
"""
Generate a GOAP-styled Word document with example ocean account tables.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# GOAP colour palette
GREEN = "3B9C7B"
TEAL = "0A5455"
TEXT_DARK = "30302F"
BORDER_CLR = "404040"
WHITE = "FFFFFF"

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "example_account_tables_print.docx")


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, colour_hex):
    """Apply background shading to a cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colour_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, colour=BORDER_CLR, sz="8"):
    """Set 1-pt solid borders on all four sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="{sz}" w:space="0" w:color="{colour}"/>'
        f'  <w:left   w:val="single" w:sz="{sz}" w:space="0" w:color="{colour}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{colour}"/>'
        f'  <w:right  w:val="single" w:sz="{sz}" w:space="0" w:color="{colour}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def style_paragraph(paragraph, font_size=10, bold=False, colour=TEXT_DARK,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Apply GOAP font styles + spacing to a paragraph."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.alignment = alignment
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(colour)


def write_cell(cell, text, font_size=10, bold=False, colour=TEXT_DARK,
               alignment=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
    """Write styled text into a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    style_paragraph(p, font_size=font_size, bold=bold, colour=colour,
                    alignment=alignment)
    set_cell_borders(cell)
    if bg:
        set_cell_shading(cell, bg)


def header_cell(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Green header cell."""
    write_cell(cell, text, bold=True, colour=WHITE, alignment=alignment,
               bg=GREEN)


def row_header_cell(cell, text):
    """Teal row-header cell."""
    write_cell(cell, text, bold=True, colour=WHITE, bg=TEAL)


def data_cell(cell, text, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    """Normal data cell."""
    write_cell(cell, text, alignment=alignment)


def add_section_title(doc, title):
    """Add a bold section heading."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


# ── Table builders ───────────────────────────────────────────────────────────

def table1_extent(doc):
    """Table 1: Extent Account (ha)."""
    add_section_title(doc, "Table 1 -- Extent Account (ha)")
    headers = ["Ecosystem type", "Opening extent\n(2018)", "Additions",
               "Reductions", "Net change", "Closing extent\n(2023)"]
    rows = [
        ["Coral reef",    "1 240", "40",  "100", "-60",  "1 180"],
        ["Seagrass beds", "860",   "35",  "70",  "-35",  "825"],
        ["Mangroves",     "315",   "40",  "15",  "+25",  "340"],
        ["Other marine",  "7 585", "120", "50",  "+70",  "7 655"],
        ["Total",         "10 000","235", "235", "0",    "10 000"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        header_cell(tbl.rows[0].cells[i], h,
                    WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c == 0:
                if row[0] == "Total":
                    write_cell(tbl.rows[r].cells[c], val, bold=True, bg=GREEN,
                               colour=WHITE)
                else:
                    row_header_cell(tbl.rows[r].cells[c], val)
            else:
                if row[0] == "Total":
                    write_cell(tbl.rows[r].cells[c], val, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER, bg=GREEN,
                               colour=WHITE)
                else:
                    data_cell(tbl.rows[r].cells[c],  val,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)


def table1b_timeseries(doc):
    """Table 1b: Extent Time Series."""
    add_section_title(doc, "Table 1b -- Extent Time-Series (ha)")
    headers = ["Ecosystem type", "2010", "2014", "2018", "2023",
               "Total change", "% change"]
    rows = [
        ["Coral reef",    "1 340", "1 290", "1 240", "1 180", "-160", "-11.9%"],
        ["Seagrass beds", "920",   "890",   "860",   "825",   "-95",  "-10.3%"],
        ["Mangroves",     "280",   "295",   "315",   "340",   "+60",  "+21.4%"],
        ["Total",         "10 000","10 000","10 000","10 000", "0",    "0.0%"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        header_cell(tbl.rows[0].cells[i], h,
                    WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c == 0:
                if row[0] == "Total":
                    write_cell(tbl.rows[r].cells[c], val, bold=True, bg=GREEN,
                               colour=WHITE)
                else:
                    row_header_cell(tbl.rows[r].cells[c], val)
            else:
                if row[0] == "Total":
                    write_cell(tbl.rows[r].cells[c], val, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER, bg=GREEN,
                               colour=WHITE)
                else:
                    data_cell(tbl.rows[r].cells[c], val,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)


def table2_change_matrix(doc):
    """Table 2: Change Matrix (ha, 2018-2023)."""
    add_section_title(doc, "Table 2 -- Ecosystem Extent Change Matrix (ha, 2018-2023)")
    headers = ["From \\ To", "Coral reef", "Seagrass", "Mangrove", "Other",
               "Opening total"]
    rows = [
        ["Coral reef",    "1 140",   "15",  "0",   "85",  "1 240"],
        ["Seagrass beds", "10",    "790",  "5",  "55",  "860"],
        ["Mangroves",     "0",     "0",   "300",  "15",  "315"],
        ["Other marine",  "30",    "20",  "35",   "7 500",  "7 585"],
        ["Closing total", "1 180",   "825",  "340",  "7 655",  "10 000"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        header_cell(tbl.rows[0].cells[i], h, WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c == 0:
                if "Total" in row[0]:
                    write_cell(tbl.rows[r].cells[c], val, bold=True, bg=GREEN,
                               colour=WHITE)
                else:
                    row_header_cell(tbl.rows[r].cells[c], val)
            else:
                if "Total" in rows[r-1][0] or "Closing" in rows[r-1][0]:
                    write_cell(tbl.rows[r].cells[c], val, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER, bg=GREEN,
                               colour=WHITE)
                elif r == c:  # Diagonal = stable area, highlight mint
                    data_cell(tbl.rows[r].cells[c], val,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    set_cell_shading(tbl.rows[r].cells[c], "D4EEE5")
                else:
                    data_cell(tbl.rows[r].cells[c], val,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)


def table3_condition(doc):
    """Table 3: Condition Account."""
    add_section_title(doc, "Table 3 -- Ecosystem Condition Account (Coral Reef)")
    headers = ["Indicator", "Unit", "Reference\nlevel",
               "Opening CI\n(2018)", "Closing CI\n(2023)",
               "Change", "Interpretation"]
    rows = [
        ["Live coral cover",      "%",         "40",  "0.75", "0.68", "-0.07", "Declining"],
        ["Macroalgae cover",      "%",         "10",  "0.60", "0.52", "-0.08", "Declining"],
        ["Fish biomass",          "kg/ha",     "500", "0.82", "0.78", "-0.04", "Declining"],
        ["Coral recruit density", "per m\u00B2","8",  "0.63", "0.55", "-0.08", "Declining"],
        ["Reef structural relief","index 1-5", "4.0", "0.70", "0.65", "-0.05", "Declining"],
        ["Bleaching prevalence",  "% colonies","5",   "0.80", "0.60", "-0.20", "Declining"],
        ["Composite Condition Index","average", "",  "0.72", "0.63", "-0.09", "Declining"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        header_cell(tbl.rows[0].cells[i], h, WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c == 0:
                if "Composite" in row[0]:
                    write_cell(tbl.rows[r].cells[c], val, bold=True, bg=GREEN,
                               colour=WHITE)
                else:
                    row_header_cell(tbl.rows[r].cells[c], val)
            else:
                if "Composite" in rows[r-1][0]:
                    write_cell(tbl.rows[r].cells[c], val, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER, bg=GREEN,
                               colour=WHITE)
                else:
                    data_cell(tbl.rows[r].cells[c], val,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)


def table4a_services_physical(doc):
    """Table 4a: Ecosystem Services -- Physical Supply."""
    add_section_title(doc, "Table 4a -- Ecosystem Services Supply (Physical Units)")
    headers = ["Service", "Unit", "Coral reef", "Seagrass", "Mangrove", "Total"]
    rows = [
        ["Fish provisioning",     "kg/yr",       "45 000",  "28 000", "18 000",  "91 000"],
        ["Carbon sequestration",  "Mg CO\u2082/yr","1 200", "3 800",  "5 600",   "10 600"],
        ["Coastal protection",    "m shoreline", "12 000",  "4 500",  "8 200",   "24 700"],
        ["Nursery habitat",       "kg recruits/yr","8 500", "12 000", "6 200",   "26 700"],
        ["Recreation & tourism",  "visitors/yr", "85 000",  "15 000", "22 000",  "122 000"],
        ["Gleaning",              "hours/yr",    "3 200",   "6 800",  "1 500",   "11 500"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        header_cell(tbl.rows[0].cells[i], h, WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c == 0:
                row_header_cell(tbl.rows[r].cells[c], val)
            else:
                data_cell(tbl.rows[r].cells[c], val,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)


def table4b_services_monetary(doc):
    """Table 4b: Ecosystem Services -- Monetary Supply."""
    add_section_title(doc, "Table 4b -- Ecosystem Services Supply (Monetary, USD/yr)")
    headers = ["Service", "Valuation method", "Coral reef", "Seagrass",
               "Mangrove", "Total"]
    rows = [
        ["Fish provisioning",    "Resource rent",      "630 000",   "336 000",
         "198 000",   "1 164 000"],
        ["Carbon sequestration", "Social cost of carbon","60 000",  "190 000",
         "280 000",   "530 000"],
        ["Coastal protection",   "Replacement cost",   "2 400 000", "540 000",
         "1 640 000", "4 580 000"],
        ["Recreation & tourism", "Expenditure method", "4 250 000", "375 000",
         "660 000",   "5 285 000"],
        ["Gleaning",             "Equivalent wage",    "19 200",    "40 800",
         "9 000",     "69 000"],
        ["Total",                "",                   "7 359 200", "1 481 800",
         "2 787 000", "11 628 000"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        header_cell(tbl.rows[0].cells[i], h, WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c == 0:
                if row[0] == "Total":
                    write_cell(tbl.rows[r].cells[c], val, bold=True, bg=GREEN,
                               colour=WHITE)
                else:
                    row_header_cell(tbl.rows[r].cells[c], val)
            else:
                if rows[r-1][0] == "Total":
                    write_cell(tbl.rows[r].cells[c], val, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER, bg=GREEN,
                               colour=WHITE)
                else:
                    data_cell(tbl.rows[r].cells[c], val,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)


def table4c_services_use_physical(doc):
    """Table 4c: Physical Use -- Who Benefits."""
    add_section_title(doc, "Table 4c -- Ecosystem Service Physical Use (Part A): Who Benefits?")
    headers = ["Service", "Unit", "Fisheries\nsector", "Tourism\nsector",
               "Coastal\nhouseholds", "Government", "Global\ncommunity", "Total"]
    rows = [
        ["Fish provisioning",    "kg/yr",      "180 000", "0",      "0",      "0",      "0",      "180 000"],
        ["Carbon sequestration", "Mg CO2/yr",  "0",       "0",      "0",      "0",      "2 335",  "2 335"],
        ["Coastal protection",   "m coastline","0",       "0",      "10 000", "5 500",  "0",      "15 500"],
        ["Nursery habitat",      "kg/yr",      "7 300",   "0",      "0",      "0",      "0",      "7 300"],
        ["Recreation",           "visitors/yr","0",       "15 000", "2 400",  "0",      "0",      "17 400"],
        ["Gleaning",             "hours/yr",   "0",       "0",      "18 000", "0",      "0",      "18 000"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        header_cell(tbl.rows[0].cells[i], h, WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c == 0:
                row_header_cell(tbl.rows[r].cells[c], val)
            else:
                data_cell(tbl.rows[r].cells[c], val, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    style_paragraph(p, font_size=9, colour="717171")
    p.add_run("Supply is organized by ecosystem. Use is organized by beneficiary. Same quantities, different perspective.").font.italic = True


def table4d_services_use_monetary(doc):
    """Table 4d: Monetary Use -- Who Benefits."""
    add_section_title(doc, "Table 4d -- Ecosystem Service Monetary Use (Part B): USD/yr by Beneficiary")
    headers = ["Service", "Method", "Fisheries", "Tourism",
               "Coastal HH", "Government", "Global", "Total"]
    rows = [
        ["Fish provisioning",    "Resource rent",     "210 000",  "0",         "0",       "0",       "0",       "210 000"],
        ["Carbon sequestration", "SCC",               "0",        "0",         "0",       "0",       "119 085", "119 085"],
        ["Coastal protection",   "Replacement cost",  "0",        "0",         "880 000", "495 000", "0",       "1 375 000"],
        ["Nursery habitat",      "Productivity change","127 750", "0",         "0",       "0",       "0",       "127 750"],
        ["Recreation",           "Expenditure",       "0",        "1 275 000", "84 000",  "0",       "0",       "1 359 000"],
        ["Gleaning",             "Equiv. wage",       "0",        "0",         "63 000",  "0",       "0",       "63 000"],
        ["Total",                "",                  "337 750",  "1 275 000", "1 027 000","495 000","119 085", "3 253 835"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        header_cell(tbl.rows[0].cells[i], h, WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c == 0:
                if row[0] == "Total":
                    write_cell(tbl.rows[r].cells[c], val, bold=True, bg=GREEN, colour=WHITE)
                else:
                    row_header_cell(tbl.rows[r].cells[c], val)
            else:
                if rows[r-1][0] == "Total":
                    write_cell(tbl.rows[r].cells[c], val, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER, bg=GREEN, colour=WHITE)
                else:
                    data_cell(tbl.rows[r].cells[c], val, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    style_paragraph(p, font_size=9, colour="717171")
    p.add_run("Tourism sector receives the most value (39%). Coastal households receive 32% (protection + gleaning + mangrove recreation).").font.italic = True


def merge_cells_styled(tbl, row_idx, start_col, end_col, text, bg_colour):
    """Merge cells in a python-docx table row and style them."""
    # Merge
    cell_start = tbl.rows[row_idx].cells[start_col]
    cell_end = tbl.rows[row_idx].cells[end_col]
    cell_start.merge(cell_end)
    # Style the merged cell
    write_cell(cell_start, text, bold=True, colour=WHITE,
               alignment=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_colour)


def _build_seea_table(doc, table_type, bg_colour):
    """Build a SEEA EA Table 7.1 supply or use sub-table.

    table_type: 'supply' or 'use'
    bg_colour: header background colour hex string
    """
    LIGHT_BG = "D5E8D4" if table_type == "supply" else "B2D8D8"

    services = ["Fish provisioning", "Carbon sequestration", "Coastal protection",
                "Nursery habitat support", "Recreation & tourism",
                "Gleaning (subsistence harvest)"]
    units = ["kg/yr", "Mg CO\u2082/yr", "m of coastline", "kg fish/yr",
             "visitors/yr", "hours/yr"]
    # Categories: list of (label, count_of_services_after)
    categories = [
        ("Provisioning services (6.1)", 1),
        ("Regulating & maintenance services (6.2\u20136.5)", 3),
        ("Cultural services (6.6)", 2),
    ]

    # Columns: Service, Unit, Fisheries, Tourism, Other, Households, Govt,
    #          Global/Exports, Coral reefs, Seagrass, Mangroves, Total eco, TOTAL
    col_headers = ["Service", "Unit", "Fisheries", "Tourism", "Other industry",
                   "Households", "Govt", "Global/\nExports",
                   "Coral reefs\n(M1.3)", "Seagrass\n(M1.1)",
                   "Mangroves\n(MFT1.2)", "Total\necosystem", "TOTAL"]
    n_cols = len(col_headers)

    if table_type == "supply":
        # col indices 2-12 (0-based from col C): only ecosystem cols have values
        # index mapping: 0=Fisheries..5=Global, 6=Coral,7=Seagrass,8=Mangroves,9=TotalEco,10=TOTAL
        data = [
            {"6": "120,000", "7": "45,000", "8": "15,000", "9": "180,000", "10": "180,000"},
            {"7": "1,040", "8": "1,295", "9": "2,335", "10": "2,335"},
            {"6": "12,000", "8": "3,500", "9": "15,500", "10": "15,500"},
            {"6": "4,500", "7": "2,800", "9": "7,300", "10": "7,300"},
            {"6": "15,000", "8": "2,400", "9": "17,400", "10": "17,400"},
            {"7": "18,000", "9": "18,000", "10": "18,000"},
        ]
    else:
        data = [
            {"0": "180,000", "10": "180,000"},
            {"5": "2,335", "10": "2,335"},
            {"3": "10,000", "4": "5,500", "10": "15,500"},
            {"0": "7,300", "10": "7,300"},
            {"1": "15,000", "3": "2,400", "10": "17,400"},
            {"3": "18,000", "10": "18,000"},
        ]

    # Count total rows: 2 header rows + 3 category rows + 6 data rows
    n_rows = 2 + 3 + len(services)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: group header -- merge A-B, C-H "Economic units", I-K "Ecosystem types", L-M blank
    for c in range(n_cols):
        write_cell(tbl.rows[0].cells[c], "", bg=bg_colour)
    merge_cells_styled(tbl, 0, 2, 7, "Economic units", bg_colour)
    merge_cells_styled(tbl, 0, 8, 10, "Ecosystem types", bg_colour)

    # Row 1: column headers
    for i, h in enumerate(col_headers):
        header_cell(tbl.rows[1].cells[i], h, WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(tbl.rows[1].cells[i], bg_colour)

    # Data rows with category sub-headers
    row_idx = 2
    svc_idx = 0
    for cat_label, cat_count in categories:
        # Category row
        merge_cells_styled(tbl, row_idx, 0, n_cols - 1, cat_label, LIGHT_BG)
        # Make category text dark
        for run in tbl.rows[row_idx].cells[0].paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string("404040")
            run.font.bold = True
            run.font.italic = True
            run.font.size = Pt(9)
        row_idx += 1

        for _ in range(cat_count):
            row_header_cell(tbl.rows[row_idx].cells[0], services[svc_idx])
            data_cell(tbl.rows[row_idx].cells[1], units[svc_idx],
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(tbl.rows[row_idx].cells[1], TEAL)
            for run in tbl.rows[row_idx].cells[1].paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string(WHITE)
                run.font.size = Pt(9)

            svc_data = data[svc_idx]
            for j in range(11):
                val = svc_data.get(str(j), "")
                data_cell(tbl.rows[row_idx].cells[2 + j], val,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

            svc_idx += 1
            row_idx += 1


def table4e_combined_supply_use(doc):
    """Table 4e: Combined Supply and Use -- SEEA EA Table 7.1 format."""
    DARK_GREEN = "2C745B"

    add_section_title(doc, "Table 4e -- Supply and Use of Ecosystem Services (SEEA EA Table 7.1)")

    # Supply table (Table 7.1a)
    p = doc.add_paragraph()
    style_paragraph(p, font_size=11, colour=GREEN)
    p.add_run("SUPPLY TABLE (Table 7.1a)").bold = True

    _build_seea_table(doc, "supply", GREEN)

    doc.add_paragraph()  # spacer

    # Use table (Table 7.1b)
    p = doc.add_paragraph()
    style_paragraph(p, font_size=11, colour=TEAL)
    p.add_run("USE TABLE (Table 7.1b)").bold = True

    _build_seea_table(doc, "use", DARK_GREEN)

    p = doc.add_paragraph()
    style_paragraph(p, font_size=9, colour="717171")
    p.add_run(
        "Format follows SEEA EA Table 7.1 (supply and use of ecosystem services). "
        "Supply table shows which ecosystems generate each service; "
        "use table shows which economic sectors benefit. "
        "Carbon sequestration use is allocated to Global/Exports."
    ).font.italic = True


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)  # landscape
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Title
    title = doc.add_paragraph()
    run = title.add_run("GOAP Example Ocean Account Tables")
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Illustrative data for a hypothetical coastal ecosystem accounting area")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Build each table
    table1_extent(doc)
    table1b_timeseries(doc)
    doc.add_page_break()
    table2_change_matrix(doc)
    table3_condition(doc)
    doc.add_page_break()
    table4a_services_physical(doc)
    table4b_services_monetary(doc)
    doc.add_page_break()
    table4c_services_use_physical(doc)
    table4d_services_use_monetary(doc)
    doc.add_page_break()
    table4e_combined_supply_use(doc)

    doc.save(OUTPUT_FILE)
    print(f"Saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
