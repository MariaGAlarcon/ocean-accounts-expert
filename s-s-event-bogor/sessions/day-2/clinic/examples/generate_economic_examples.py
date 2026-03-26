#!/usr/bin/env python3
"""
Generate GOAP-styled Excel and Word files with example OESA, OTSA,
and Waste account tables for ocean accounting.
"""

import os
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── GOAP palette ────────────────────────────────────────────────────────────
GREEN = "3B9C7B"
TEAL = "0A5455"
TEXT_CLR = "30302F"
GRAY = "404040"
WHITE = "FFFFFF"

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
XLSX_FILE = os.path.join(OUTPUT_DIR, "example_economic_accounts.xlsx")
DOCX_FILE = os.path.join(OUTPUT_DIR, "example_economic_accounts_print.docx")

# ── openpyxl style objects ──────────────────────────────────────────────────
FONT_HEADER = Font(name="Arial", size=10, bold=True, color=WHITE)
FONT_ROW_HDR = Font(name="Arial", size=10, bold=True, color=WHITE)
FONT_DATA = Font(name="Arial", size=10, color=TEXT_CLR)
FONT_TOTAL = Font(name="Arial", size=10, bold=True, color=WHITE)
FONT_TITLE = Font(name="Arial", size=13, bold=True, color=TEAL)
FONT_SUBTITLE = Font(name="Arial", size=10, italic=True, color=TEXT_CLR)

FILL_GREEN = PatternFill(start_color=GREEN, end_color=GREEN, fill_type="solid")
FILL_TEAL = PatternFill(start_color=TEAL, end_color=TEAL, fill_type="solid")
FILL_NONE = PatternFill(fill_type=None)

THIN_BORDER = Border(
    left=Side(style="thin", color=GRAY),
    right=Side(style="thin", color=GRAY),
    top=Side(style="thin", color=GRAY),
    bottom=Side(style="thin", color=GRAY),
)

ALIGN_LEFT = Alignment(horizontal="left", vertical="center", wrap_text=True)
ALIGN_CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
ALIGN_RIGHT = Alignment(horizontal="right", vertical="center", wrap_text=True)

NUM_FMT = '#,##0'
PCT_FMT = '0.0%'


# ── Excel helpers ───────────────────────────────────────────────────────────

def xl_header_row(ws, row, values, col_start=1):
    for c, val in enumerate(values, start=col_start):
        cell = ws.cell(row=row, column=c, value=val)
        cell.font = FONT_HEADER
        cell.fill = FILL_GREEN
        cell.border = THIN_BORDER
        cell.alignment = ALIGN_CENTER


def xl_data_row(ws, row, values, col_start=1, is_total=False, num_cols=None,
                pct_cols=None, text_cols=None):
    """Write a data row. First column is row header (teal) unless is_total."""
    if num_cols is None:
        num_cols = set()
    if pct_cols is None:
        pct_cols = set()
    if text_cols is None:
        text_cols = set()
    for c, val in enumerate(values, start=col_start):
        cell = ws.cell(row=row, column=c)
        cell.border = THIN_BORDER
        if is_total:
            cell.font = FONT_TOTAL
            cell.fill = FILL_GREEN
            cell.alignment = ALIGN_CENTER if c > col_start else ALIGN_LEFT
            cell.value = val
        elif c == col_start:
            # row header
            cell.font = FONT_ROW_HDR
            cell.fill = FILL_TEAL
            cell.alignment = ALIGN_LEFT
            cell.value = val
        else:
            cell.font = FONT_DATA
            cell.fill = FILL_NONE
            cell.alignment = ALIGN_RIGHT
            cell.value = val
        # number formats
        if c in num_cols and isinstance(val, (int, float)):
            cell.number_format = NUM_FMT
        if c in pct_cols and isinstance(val, (int, float)):
            cell.number_format = PCT_FMT


def xl_title(ws, row, text, col_span=5):
    cell = ws.cell(row=row, column=1, value=text)
    cell.font = FONT_TITLE
    cell.alignment = ALIGN_LEFT
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_span)


def xl_subtitle(ws, row, text, col_span=5):
    cell = ws.cell(row=row, column=1, value=text)
    cell.font = FONT_SUBTITLE
    cell.alignment = ALIGN_LEFT
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_span)


def auto_width(ws, min_width=12, max_width=30):
    from openpyxl.cell.cell import MergedCell
    for col in ws.columns:
        length = min_width
        first_cell = None
        for cell in col:
            if isinstance(cell, MergedCell):
                continue
            if first_cell is None:
                first_cell = cell
            if cell.value:
                length = max(length, min(len(str(cell.value)) + 4, max_width))
        if first_cell is not None:
            ws.column_dimensions[first_cell.column_letter].width = length


# ── Sheet builders (Excel) ──────────────────────────────────────────────────

def sheet_oesa_output(wb):
    ws = wb.create_sheet("OESA Output")
    cols = 2
    xl_title(ws, 1, "Table 5: Ocean Economy Gross Output by Sector", cols)
    xl_subtitle(ws, 2, "Example: South Africa 2018 (R million)", cols)
    headers = ["Ocean Economy Group", "Gross Output (R million)"]
    xl_header_row(ws, 4, headers)
    data = [
        ("Living Resources", 20256),
        ("Marine Minerals", 142769),
        ("Marine Construction", 135414),
        ("Ship and Boat Building", 5555),
        ("Marine Transport", 43430),
        ("Coastal Tourism", 45841),
    ]
    r = 5
    nc = {2}
    for name, val in data:
        xl_data_row(ws, r, [name, val], num_cols=nc)
        r += 1
    xl_data_row(ws, r, ["Total", 393265], is_total=True, num_cols=nc)
    auto_width(ws)


def sheet_oesa_ic(wb):
    ws = wb.create_sheet("OESA IC")
    cols = 2
    xl_title(ws, 1, "Table 6: Ocean Economy Intermediate Consumption by Sector", cols)
    xl_subtitle(ws, 2, "Example: South Africa 2018 (R million)", cols)
    headers = ["Ocean Economy Group", "IC (R million)"]
    xl_header_row(ws, 4, headers)
    data = [
        ("Living Resources", 9772),
        ("Marine Minerals", 96859),
        ("Marine Construction", 94802),
        ("Ship and Boat Building", 4975),
        ("Marine Transport", 27652),
        ("Coastal Tourism", 26711),
    ]
    r = 5
    nc = {2}
    for name, val in data:
        xl_data_row(ws, r, [name, val], num_cols=nc)
        r += 1
    xl_data_row(ws, r, ["Total", 260771], is_total=True, num_cols=nc)
    auto_width(ws)


def sheet_oesa_gva(wb):
    ws = wb.create_sheet("OESA GVA")
    cols = 4
    xl_title(ws, 1, "Table 7: Ocean Economy Gross Value Added", cols)
    xl_subtitle(ws, 2, "Example: South Africa 2018", cols)
    headers = ["Ocean Economy Group", "GVA (R million)", "Share of Ocean GVA (%)", "Employment"]
    xl_header_row(ws, 4, headers)
    data = [
        ("Living Resources", 10485, 0.079, 27000),
        ("Marine Minerals", 45910, 0.347, 48000),
        ("Marine Construction", 40612, 0.307, 125000),
        ("Ship and Boat Building", 580, 0.004, 4500),
        ("Marine Transport", 15778, 0.119, 32000),
        ("Coastal Tourism", 19130, 0.144, 95000),
    ]
    r = 5
    nc = {2, 4}
    pc = {3}
    for row in data:
        xl_data_row(ws, r, list(row), num_cols=nc, pct_cols=pc)
        r += 1
    xl_data_row(ws, r, ["Total", 132495, 1.0, 331500], is_total=True, num_cols=nc, pct_cols=pc)
    # fix total pct format
    ws.cell(row=r, column=3).number_format = '0%'
    auto_width(ws)


def sheet_oesa_summary(wb):
    ws = wb.create_sheet("OESA Summary")
    cols = 2
    xl_title(ws, 1, "Table 8: Ocean Economy Summary", cols)
    xl_subtitle(ws, 2, "Example: South Africa 2018", cols)
    headers = ["Indicator", "Value"]
    xl_header_row(ws, 4, headers)
    data = [
        ("Total Ocean Economy GVA", "R 132,495 million"),
        ("National GDP", "R 4,829,603 million"),
        ("Ocean Economy % of GDP", "2.74%"),
        ("Ocean Economy Employment", "331,500"),
    ]
    r = 5
    for name, val in data:
        xl_data_row(ws, r, [name, val], text_cols={2})
        r += 1
    auto_width(ws)


def sheet_otsa(wb):
    ws = wb.create_sheet("OTSA")
    cols = 4
    xl_title(ws, 1, "Table 9: Ocean Tourism Satellite Account", cols)
    xl_subtitle(ws, 2, "Example: Illustrative coastal SIDS", cols)
    headers = ["Indicator", "Total Tourism", "Tourism Partial", "Coastal Tourism"]
    xl_header_row(ws, 4, headers)
    data = [
        ("International arrivals", 500000, "35%", 175000),
        ("Domestic trips", 2000000, "20%", 400000),
        ("Tourism expenditure (USD million)", 800, "25.8%", 206),
        ("Tourism direct GVA (USD million)", 320, "25.8%", 82.5),
        ("Tourism employment", 45000, "25.8%", 11610),
        ("National GDP (USD million)", 5000, "", ""),
        ("Coastal tourism % of GDP", "", "", "1.65%"),
    ]
    r = 5
    nc = {2, 4}
    for row in data:
        xl_data_row(ws, r, list(row), num_cols=nc)
        r += 1
    auto_width(ws)


def sheet_water_emissions(wb):
    ws = wb.create_sheet("Water Emissions")
    cols = 5
    xl_title(ws, 1, "Table 10: Water Emissions Account - Substances Discharged to Sea", cols)
    xl_subtitle(ws, 2, "Supply side: generation by source", cols)
    headers = ["Source", "Nitrogen (tonnes N/yr)", "Phosphorus (tonnes P/yr)",
               "BOD (tonnes/yr)", "Heavy metals (tonnes/yr)"]
    xl_header_row(ws, 4, headers)
    data = [
        ("Agriculture runoff", 580, 95, 1200, 2),
        ("Manufacturing", 180, 35, 800, 15),
        ("Sewerage (treated)", 120, 40, 300, 1),
        ("Sewerage (untreated)", 280, 55, 1500, 5),
        ("Households (direct)", 40, 15, 200, 0),
    ]
    r = 5
    nc = {2, 3, 4, 5}
    for row in data:
        xl_data_row(ws, r, list(row), num_cols=nc)
        r += 1
    xl_data_row(ws, r, ["Total", 1200, 240, 4000, 23], is_total=True, num_cols=nc)
    auto_width(ws)


def sheet_solid_waste(wb):
    ws = wb.create_sheet("Solid Waste")
    cols = 5
    xl_title(ws, 1, "Table 11: Solid Waste Entering the Marine Environment", cols)
    xl_subtitle(ws, 2, "", cols)
    headers = ["Waste type", "Land-based (tonnes/yr)", "Sea-based (tonnes/yr)",
               "Total (tonnes/yr)", "% of total"]
    xl_header_row(ws, 4, headers)
    data = [
        ("Plastics (macro)", 18000, 1200, 19200, "52%"),
        ("Plastics (micro)", 3500, 500, 4000, "11%"),
        ("Fishing gear", 200, 3800, 4000, "11%"),
        ("Ship waste", 0, 2500, 2500, "7%"),
        ("Sewage sludge", 4000, 0, 4000, "11%"),
        ("Other", 2300, 1000, 3300, "9%"),
    ]
    r = 5
    nc = {2, 3, 4}
    for row in data:
        xl_data_row(ws, r, list(row), num_cols=nc)
        r += 1
    xl_data_row(ws, r, ["Total", 28000, 9000, 37000, "100%"], is_total=True, num_cols=nc)
    auto_width(ws)


def sheet_air_emissions(wb):
    ws = wb.create_sheet("Air Emissions")
    cols = 5
    xl_title(ws, 1, "Table 12: Air Emissions from Ocean Industries", cols)
    xl_subtitle(ws, 2, "", cols)
    headers = ["Ocean sector", "Fuel (tonnes/yr)", "CO2 (tonnes/yr)",
               "SO2 (tonnes/yr)", "NOx (tonnes/yr)"]
    xl_header_row(ws, 4, headers)
    data = [
        ("International shipping", 150000, 480000, 8100, 12600),
        ("Domestic shipping", 50000, 160000, 2700, 4200),
        ("Fishing fleet", 50000, 160000, 2700, 4200),
        ("Offshore oil and gas", 30000, 96000, 450, 1800),
        ("Port operations", 20000, 64000, 300, 1200),
    ]
    r = 5
    nc = {2, 3, 4, 5}
    for row in data:
        xl_data_row(ws, r, list(row), num_cols=nc)
        r += 1
    xl_data_row(ws, r, ["Total", 300000, 960000, 14250, 24000], is_total=True, num_cols=nc)
    auto_width(ws)


def sheet_summary_all(wb):
    ws = wb.create_sheet("Summary All Accounts")
    cols = 4
    xl_title(ws, 1, "Complete Ocean Accounting Framework: All Account Types", cols)
    xl_subtitle(ws, 2, "", cols)
    headers = ["Account Type", "Framework", "Key Indicator", "Example Value"]
    xl_header_row(ws, 4, headers)
    data = [
        ("Extent", "SEEA-EA", "Ecosystem area", "10,000 ha (3 ecosystems)"),
        ("Condition", "SEEA-EA", "Composite CI", "0.53 (declining)"),
        ("Ecosystem Services (physical)", "SEEA-EA", "9 services quantified", "Multiple units"),
        ("Ecosystem Services (monetary)", "SEEA-EA", "Total service value", "USD 3.7 million/yr"),
        ("Ocean Economy (OESA)", "SNA", "Ocean GDP", "R 132.5 billion (2.74%)"),
        ("Ocean Tourism (OTSA)", "SNA/TSA", "Coastal tourism GVA", "USD 82.5 million (1.65%)"),
        ("Waste - Water emissions", "SEEA-CF", "Nutrient load to sea", "1,200 t N + 240 t P/yr"),
        ("Waste - Solid waste", "SEEA-CF", "Marine waste", "37,000 tonnes/yr"),
        ("Waste - Air emissions", "SEEA-CF", "Ocean industry CO2", "960,000 tonnes CO2/yr"),
    ]
    r = 5
    for row in data:
        xl_data_row(ws, r, list(row), text_cols={2, 3, 4})
        r += 1
    auto_width(ws, min_width=14, max_width=35)


# ── Word doc helpers (matching generate_printable_examples.py) ──────────────

def set_cell_shading(cell, colour_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colour_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, colour=GRAY, sz="8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="{sz}" w:space="0" w:color="{colour}"/>'
        f'  <w:left   w:val="single" w:sz="{sz}" w:space="0" w:color="{colour}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{colour}"/>'
        f'  <w:right  w:val="single" w:sz="{sz}" w:space="0" w:color="{colour}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def style_paragraph(paragraph, font_size=10, bold=False, colour=TEXT_CLR,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.alignment = alignment
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(colour)


def write_cell(cell, text, font_size=10, bold=False, colour=TEXT_CLR,
               alignment=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(str(text))
    style_paragraph(p, font_size=font_size, bold=bold, colour=colour,
                    alignment=alignment)
    set_cell_borders(cell)
    if bg:
        set_cell_shading(cell, bg)


def doc_header_cell(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    write_cell(cell, text, bold=True, colour=WHITE, alignment=alignment, bg=GREEN)


def doc_row_header_cell(cell, text):
    write_cell(cell, text, bold=True, colour=WHITE, bg=TEAL)


def doc_data_cell(cell, text, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    write_cell(cell, text, alignment=alignment)


def doc_total_cell(cell, text, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    write_cell(cell, text, bold=True, colour=WHITE, alignment=alignment, bg=GREEN)


def add_section_title(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(TEXT_CLR)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def build_doc_table(doc, title, subtitle, headers, rows, total_row=None):
    """Generic doc table builder."""
    add_section_title(doc, title)
    if subtitle:
        add_subtitle(doc, subtitle)
    n_rows = len(rows) + 1 + (1 if total_row else 0)
    n_cols = len(headers)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # headers
    for i, h in enumerate(headers):
        doc_header_cell(tbl.rows[0].cells[i], h,
                        WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)
    # data rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            if c_idx == 0:
                doc_row_header_cell(tbl.rows[r_idx].cells[c_idx], str(val))
            else:
                doc_data_cell(tbl.rows[r_idx].cells[c_idx], str(val),
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # total row
    if total_row:
        tr = len(rows) + 1
        for c_idx, val in enumerate(total_row):
            doc_total_cell(tbl.rows[tr].cells[c_idx], str(val),
                           WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT)


# ── Word doc table builders ─────────────────────────────────────────────────

def doc_oesa_output(doc):
    build_doc_table(
        doc,
        "Table 5: Ocean Economy Gross Output by Sector",
        "Example: South Africa 2018 (R million)",
        ["Ocean Economy Group", "Gross Output (R million)"],
        [
            ["Living Resources", "20,256"],
            ["Marine Minerals", "142,769"],
            ["Marine Construction", "135,414"],
            ["Ship and Boat Building", "5,555"],
            ["Marine Transport", "43,430"],
            ["Coastal Tourism", "45,841"],
        ],
        total_row=["Total", "393,265"],
    )


def doc_oesa_ic(doc):
    build_doc_table(
        doc,
        "Table 6: Ocean Economy Intermediate Consumption by Sector",
        "Example: South Africa 2018 (R million)",
        ["Ocean Economy Group", "IC (R million)"],
        [
            ["Living Resources", "9,772"],
            ["Marine Minerals", "96,859"],
            ["Marine Construction", "94,802"],
            ["Ship and Boat Building", "4,975"],
            ["Marine Transport", "27,652"],
            ["Coastal Tourism", "26,711"],
        ],
        total_row=["Total", "260,771"],
    )


def doc_oesa_gva(doc):
    build_doc_table(
        doc,
        "Table 7: Ocean Economy Gross Value Added",
        "Example: South Africa 2018",
        ["Ocean Economy Group", "GVA (R million)", "Share of Ocean GVA (%)", "Employment"],
        [
            ["Living Resources", "10,485", "7.9%", "27,000"],
            ["Marine Minerals", "45,910", "34.7%", "48,000"],
            ["Marine Construction", "40,612", "30.7%", "125,000"],
            ["Ship and Boat Building", "580", "0.4%", "4,500"],
            ["Marine Transport", "15,778", "11.9%", "32,000"],
            ["Coastal Tourism", "19,130", "14.4%", "95,000"],
        ],
        total_row=["Total", "132,495", "100%", "331,500"],
    )


def doc_oesa_summary(doc):
    build_doc_table(
        doc,
        "Table 8: Ocean Economy Summary",
        "Example: South Africa 2018",
        ["Indicator", "Value"],
        [
            ["Total Ocean Economy GVA", "R 132,495 million"],
            ["National GDP", "R 4,829,603 million"],
            ["Ocean Economy % of GDP", "2.74%"],
            ["Ocean Economy Employment", "331,500"],
        ],
    )


def doc_otsa(doc):
    build_doc_table(
        doc,
        "Table 9: Ocean Tourism Satellite Account",
        "Example: Illustrative coastal SIDS",
        ["Indicator", "Total Tourism", "Tourism Partial", "Coastal Tourism"],
        [
            ["International arrivals", "500,000", "35%", "175,000"],
            ["Domestic trips", "2,000,000", "20%", "400,000"],
            ["Tourism expenditure (USD million)", "800", "25.8%", "206"],
            ["Tourism direct GVA (USD million)", "320", "25.8%", "82.5"],
            ["Tourism employment", "45,000", "25.8%", "11,610"],
            ["National GDP (USD million)", "5,000", "", ""],
            ["Coastal tourism % of GDP", "", "", "1.65%"],
        ],
    )


def doc_water_emissions(doc):
    build_doc_table(
        doc,
        "Table 10: Water Emissions Account - Substances Discharged to Sea",
        "Supply side: generation by source",
        ["Source", "Nitrogen (tonnes N/yr)", "Phosphorus (tonnes P/yr)",
         "BOD (tonnes/yr)", "Heavy metals (tonnes/yr)"],
        [
            ["Agriculture runoff", "580", "95", "1,200", "2"],
            ["Manufacturing", "180", "35", "800", "15"],
            ["Sewerage (treated)", "120", "40", "300", "1"],
            ["Sewerage (untreated)", "280", "55", "1,500", "5"],
            ["Households (direct)", "40", "15", "200", "0"],
        ],
        total_row=["Total", "1,200", "240", "4,000", "23"],
    )


def doc_solid_waste(doc):
    build_doc_table(
        doc,
        "Table 11: Solid Waste Entering the Marine Environment",
        None,
        ["Waste type", "Land-based (tonnes/yr)", "Sea-based (tonnes/yr)",
         "Total (tonnes/yr)", "% of total"],
        [
            ["Plastics (macro)", "18,000", "1,200", "19,200", "52%"],
            ["Plastics (micro)", "3,500", "500", "4,000", "11%"],
            ["Fishing gear", "200", "3,800", "4,000", "11%"],
            ["Ship waste", "0", "2,500", "2,500", "7%"],
            ["Sewage sludge", "4,000", "0", "4,000", "11%"],
            ["Other", "2,300", "1,000", "3,300", "9%"],
        ],
        total_row=["Total", "28,000", "9,000", "37,000", "100%"],
    )


def doc_air_emissions(doc):
    build_doc_table(
        doc,
        "Table 12: Air Emissions from Ocean Industries",
        None,
        ["Ocean sector", "Fuel (tonnes/yr)", "CO2 (tonnes/yr)",
         "SO2 (tonnes/yr)", "NOx (tonnes/yr)"],
        [
            ["International shipping", "150,000", "480,000", "8,100", "12,600"],
            ["Domestic shipping", "50,000", "160,000", "2,700", "4,200"],
            ["Fishing fleet", "50,000", "160,000", "2,700", "4,200"],
            ["Offshore oil and gas", "30,000", "96,000", "450", "1,800"],
            ["Port operations", "20,000", "64,000", "300", "1,200"],
        ],
        total_row=["Total", "300,000", "960,000", "14,250", "24,000"],
    )


def doc_summary_all(doc):
    build_doc_table(
        doc,
        "Complete Ocean Accounting Framework: All Account Types",
        None,
        ["Account Type", "Framework", "Key Indicator", "Example Value"],
        [
            ["Extent", "SEEA-EA", "Ecosystem area", "10,000 ha (3 ecosystems)"],
            ["Condition", "SEEA-EA", "Composite CI", "0.53 (declining)"],
            ["Ecosystem Services (physical)", "SEEA-EA", "9 services quantified", "Multiple units"],
            ["Ecosystem Services (monetary)", "SEEA-EA", "Total service value", "USD 3.7 million/yr"],
            ["Ocean Economy (OESA)", "SNA", "Ocean GDP", "R 132.5 billion (2.74%)"],
            ["Ocean Tourism (OTSA)", "SNA/TSA", "Coastal tourism GVA", "USD 82.5 million (1.65%)"],
            ["Waste - Water emissions", "SEEA-CF", "Nutrient load to sea", "1,200 t N + 240 t P/yr"],
            ["Waste - Solid waste", "SEEA-CF", "Marine waste", "37,000 tonnes/yr"],
            ["Waste - Air emissions", "SEEA-CF", "Ocean industry CO2", "960,000 tonnes CO2/yr"],
        ],
    )


# ── Main ────────────────────────────────────────────────────────────────────

def generate_xlsx():
    wb = Workbook()
    # remove default sheet
    wb.remove(wb.active)

    sheet_oesa_output(wb)
    sheet_oesa_ic(wb)
    sheet_oesa_gva(wb)
    sheet_oesa_summary(wb)
    sheet_otsa(wb)
    sheet_water_emissions(wb)
    sheet_solid_waste(wb)
    sheet_air_emissions(wb)
    sheet_summary_all(wb)

    wb.save(XLSX_FILE)
    print(f"Saved: {XLSX_FILE}")


def generate_docx():
    doc = Document()

    # Page setup - landscape
    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Title page
    title = doc.add_paragraph()
    run = title.add_run("GOAP Example: Economic & Waste Account Tables")
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        "Ocean Economy Satellite Account (OESA), Ocean Tourism Satellite Account (OTSA), "
        "and Waste Accounts"
    )
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(TEXT_CLR)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Build tables
    doc_oesa_output(doc)
    doc_oesa_ic(doc)
    doc.add_page_break()
    doc_oesa_gva(doc)
    doc_oesa_summary(doc)
    doc.add_page_break()
    doc_otsa(doc)
    doc_water_emissions(doc)
    doc.add_page_break()
    doc_solid_waste(doc)
    doc_air_emissions(doc)
    doc.add_page_break()
    doc_summary_all(doc)

    doc.save(DOCX_FILE)
    print(f"Saved: {DOCX_FILE}")


def main():
    generate_xlsx()
    generate_docx()
    print("Done.")


if __name__ == "__main__":
    main()
