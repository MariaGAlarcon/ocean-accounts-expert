#!/usr/bin/env python3
"""
Improved markdown-to-GOAP-docx converter.

Wraps the belize md_to_docx.py parser with fixes for:
- Horizontal rules (---) stripped instead of appearing as text
- Code fences (```) stripped, content treated as normal text
- Bullet lists (- item) rendered as proper bullet paragraphs
- Numbered lists (1. item) rendered as proper numbered paragraphs
- Inline code (`text`) backticks stripped
- Speaker notes (*Speaker notes:*) rendered as italic captions
"""

import re
import sys
from pathlib import Path

# Add belize project to path for imports
BELIZE = Path("/Users/mariaalarcon/GitHub/ocean-accounts-expert/belize-project")
sys.path.insert(0, str(BELIZE))

from md_to_docx import (
    parse_markdown as _original_parse,
    build_docx,
    DEFAULT_TEMPLATE,
    BODY_TEXT,
    GRAY_CAPTION,
    add_formatted_runs,
)
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def clean_markdown(text):
    """Pre-process markdown to fix common issues before parsing."""
    lines = text.split("\n")
    cleaned = []
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # Skip code fence markers
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        # Horizontal rules: skip entirely
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            continue

        # Clean inline backticks
        line = re.sub(r"`([^`]+)`", r"\1", line)

        # Convert bullet lists to a parseable format (prefix with special marker)
        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bullet_match and not in_code_block:
            indent = len(bullet_match.group(1))
            text_content = bullet_match.group(2)
            # Use a marker the parser will pick up
            cleaned.append(f"BULLET_{indent}: {text_content}")
            continue

        # Convert numbered lists
        num_match = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if num_match and not in_code_block:
            text_content = num_match.group(2)
            cleaned.append(f"NUMLIST: {text_content}")
            continue

        cleaned.append(line)

    return "\n".join(cleaned)


def parse_markdown_improved(text):
    """Parse markdown with improved handling of lists, code, and rules."""
    cleaned = clean_markdown(text)
    lines = cleaned.split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            blocks.append({
                "type": "heading",
                "level": len(m.group(1)),
                "text": m.group(2).strip(),
            })
            i += 1
            continue

        # Bullet list item
        if stripped.startswith("BULLET_"):
            m = re.match(r"BULLET_(\d+): (.*)", stripped)
            if m:
                blocks.append({
                    "type": "bullet",
                    "indent": int(m.group(1)),
                    "text": m.group(2),
                })
                i += 1
                continue

        # Numbered list item
        if stripped.startswith("NUMLIST: "):
            blocks.append({
                "type": "numlist",
                "text": stripped[9:],
            })
            i += 1
            continue

        # Table
        if "|" in line and stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(_parse_table_block(table_lines))
            continue

        # Caption/note lines
        if stripped and (
            stripped.lower().startswith("figure")
            or stripped.lower().startswith("table ")
            or stripped.lower().startswith("source:")
            or stripped.lower().startswith("note:")
            or stripped.startswith("*Speaker notes:")
        ):
            # Strip italic markers
            clean_text = stripped.strip("*")
            blocks.append({"type": "caption", "text": clean_text})
            i += 1
            continue

        # Blank line
        if not stripped:
            blocks.append({"type": "blank"})
            i += 1
            continue

        # Regular paragraph
        para_lines = []
        while i < len(lines):
            l = lines[i]
            ls = l.strip()
            if not ls:
                break
            if re.match(r"^#{1,6}\s+", l):
                break
            if "|" in l and ls.startswith("|"):
                break
            if ls.startswith("BULLET_") or ls.startswith("NUMLIST:"):
                break
            para_lines.append(ls)
            i += 1
        if para_lines:
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
        continue

    return blocks


def _parse_table_block(lines):
    """Parse pipe-delimited markdown table."""
    def split_row(line):
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    headers = split_row(lines[0])
    rows = []
    for line in lines[1:]:
        cells = split_row(line)
        if all(re.match(r"^[-:]+$", c) for c in cells if c):
            continue
        rows.append(cells)
    return {"type": "table", "headers": headers, "rows": rows}


def build_docx_improved(blocks, template_path, output_path, header_text=None):
    """Build docx with improved block type handling (bullets, numbered lists)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # Import formatting functions from belize converter
    from md_to_docx import format_table, remove_outer_borders, WHITE

    doc = Document(str(template_path))

    # Update headers
    if header_text:
        for section in doc.sections:
            for hdr in (section.header, section.first_page_header, section.even_page_header):
                if hdr is None:
                    continue
                for para in hdr.paragraphs:
                    if para.text.strip():
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = header_text
                        else:
                            run = para.add_run(header_text)
                            run.font.name = "Arial"
                            run.font.size = Pt(10)
                            run.font.color.rgb = BODY_TEXT

    # Clear template content
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)
    for table in doc.tables:
        table._tbl.getparent().remove(table._tbl)

    heading_styles = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3",
                      4: "Heading 4", 5: "Heading 5", 6: "Heading 6"}

    for block in blocks:
        btype = block["type"]

        if btype == "blank":
            continue

        elif btype == "heading":
            level = block["level"]
            para = doc.add_paragraph(style=heading_styles.get(level, "Heading 4"))
            add_formatted_runs(para, block["text"])

        elif btype == "paragraph":
            para = doc.add_paragraph(style="Normal")
            add_formatted_runs(para, block["text"], default_color=BODY_TEXT)

        elif btype == "bullet":
            try:
                para = doc.add_paragraph(style="List Bullet")
            except KeyError:
                para = doc.add_paragraph(style="Normal")
                para.paragraph_format.left_indent = Pt(18)
            add_formatted_runs(para, block["text"], default_color=BODY_TEXT)
            # Nested bullets
            if block.get("indent", 0) > 0:
                para.paragraph_format.left_indent = Pt(36)

        elif btype == "numlist":
            try:
                para = doc.add_paragraph(style="List Number")
            except KeyError:
                para = doc.add_paragraph(style="Normal")
                para.paragraph_format.left_indent = Pt(18)
            add_formatted_runs(para, block["text"], default_color=BODY_TEXT)

        elif btype == "caption":
            try:
                para = doc.add_paragraph(style="Figure")
            except KeyError:
                para = doc.add_paragraph(style="Normal")
            run = para.add_run(block["text"])
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = GRAY_CAPTION
            run.font.name = "Arial"

        elif btype == "table":
            headers = block["headers"]
            rows = block["rows"]
            num_cols = len(headers)

            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(h)
                run.font.name = "Arial"
                run.font.size = Pt(10)

            for r_idx, row_data in enumerate(rows):
                for j, val in enumerate(row_data):
                    if j < num_cols:
                        cell = table.rows[r_idx + 1].cells[j]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        run = para.add_run(val)
                        run.font.name = "Arial"
                        run.font.size = Pt(10)

            format_table(table, headers, rows)
            remove_outer_borders(table)

    doc.save(str(output_path))
    return output_path


def convert_file(input_path, output_path, template_path=None, header=None):
    """Convert a markdown file to GOAP-styled docx."""
    if template_path is None:
        template_path = DEFAULT_TEMPLATE

    text = Path(input_path).read_text(encoding="utf-8")
    blocks = parse_markdown_improved(text)

    if header is None:
        for b in blocks:
            if b["type"] == "heading" and b.get("level") == 1:
                header = b["text"]
                break

    build_docx_improved(blocks, template_path, output_path, header_text=header)
    print(f"  Created: {output_path}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Convert MD to GOAP Word (improved)")
    parser.add_argument("input", help="Input markdown file")
    parser.add_argument("-o", "--output", help="Output docx path")
    parser.add_argument("-t", "--template", default=str(DEFAULT_TEMPLATE))
    parser.add_argument("--header", help="Header text")
    args = parser.parse_args()

    output = args.output or str(Path(args.input).with_suffix(".docx"))
    convert_file(args.input, output, args.template, args.header)
