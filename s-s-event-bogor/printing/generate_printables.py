#!/usr/bin/env python3
"""Generate all print-ready GOAP-styled Word documents for the event."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---- GOAP Colors ----
DARK_TEAL = RGBColor(0x0A, 0x54, 0x55)
GREEN = RGBColor(0x3B, 0x9C, 0x7B)
BODY = RGBColor(0x30, 0x30, 0x2F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x71, 0x71, 0x71)
YELLOW_HEX = "FFF2CC"
MINT_HEX = "D4EEE5"


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:left w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:right w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="404040"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="404040"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def styled_run(paragraph, text, size=10, bold=False, color=BODY, font_name="Arial"):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def add_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styled_run(p, text, size=size, bold=False, color=DARK_TEAL)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    styled_run(p, text, size=11, color=GREEN)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    styled_run(p, text, size=10, color=BODY)
    return p


def add_heading_h2(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    styled_run(p, text, size=14, color=DARK_TEAL)
    return p


def add_heading_h3(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    styled_run(p, text, size=12, bold=True, color=DARK_TEAL)
    return p


def make_goap_table(doc, headers, rows, first_col_teal=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, h, size=10, bold=True, color=WHITE)
        set_cell_shading(cell, "3B9C7B")

    # Data rows
    for r, row_data in enumerate(rows, 1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            text = str(val) if val is not None else ""

            if c == 0 and first_col_teal:
                styled_run(p, text, size=10, bold=True, color=WHITE)
                set_cell_shading(cell, "0A5455")
            else:
                styled_run(p, text, size=10, color=BODY)

    set_table_borders(table)
    return table


# =========================================================================
# 1. STATION SIGNS (5 pages, landscape A4)
# =========================================================================
def generate_station_signs(outdir):
    stations = [
        ("STATION 1", "Marine Spatial Planning & MPAs",
         "Zoning, cumulative impact assessment, protected area management, spatial data",
         "0077B6"),
        ("STATION 2", "Blue Carbon & Climate",
         "Mangroves, seagrass, carbon sequestration, NDCs, climate finance, NbS",
         "2D6A4F"),
        ("STATION 3", "Fisheries & Ocean Economy",
         "Fish stocks, value chains, ocean GDP, satellite accounts, blue economy",
         "E76F51"),
        ("STATION 4", "Ecosystem Condition & Biodiversity",
         "Coral reefs, water quality, biodiversity indices, condition indicators, monitoring",
         "7B2D8E"),
        ("STATION 5", "Social Accounts & Community Engagement",
         "Livelihoods, food security, distributional equity, governance, traditional knowledge",
         "D4A017"),
    ]

    doc = Document()
    # Set landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

    for i, (label, title, subtitle, color_hex) in enumerate(stations):
        if i > 0:
            doc.add_page_break()

        # Spacer
        for _ in range(3):
            doc.add_paragraph()

        color = RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, label, size=18, color=color)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, title, size=44, bold=True, color=color)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, subtitle, size=14, color=BODY)

    path = outdir / "station_signs.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =========================================================================
# 2. CLINIC TRIAGE CARD (A4, for facilitators)
# =========================================================================
def generate_triage_card(outdir):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    add_title(doc, "Ocean Accounting Clinic -- Facilitator Triage Card", size=16)
    add_subtitle(doc, "S-S Ocean Accounting Exchange 2026, Bogor")

    add_heading_h2(doc, "Step 1: Ask three questions")

    add_body(doc, "1. What account type do you want to build?")
    for opt in [
        "   - Extent, Condition, Services (SEEA-EA ecological)",
        "   - Ocean Economy or Tourism (SNA economic)",
        "   - Waste and Emissions (SEEA-CF environmental flows)",
        "   - Social and Governance (GOAP social dimensions)",
        "   - I don't know yet -> diagnostic tool",
    ]:
        add_body(doc, opt)
    for q in [
        "2. What data do you have? (None / Global datasets / National data / Field data)",
        "3. What is your policy question?",
    ]:
        add_body(doc, q)

    add_heading_h2(doc, "Step 2: Direct to the right pathway")

    make_goap_table(doc,
        ["They want...", "They have...", "Pathway", "Materials"],
        [
            ["Extent", "No data", "A: Global datasets", "extent_account_template.xlsx"],
            ["Extent", "Imagery or maps", "B: Classification", "extent_account_template.xlsx"],
            ["Condition", "No field data", "A: RS proxies", "condition_account_template.xlsx"],
            ["Condition", "Field surveys", "B: Indicators", "condition_account_template.xlsx\ncondition_exercise.xlsx"],
            ["Services", "No data", "A: Value transfer", "services_account_template.xlsx"],
            ["Services", "Catch, visitors, etc.", "B: Primary calc", "services_account_template.xlsx\nservices_exercise.xlsx"],
            ["Ocean economy (GDP)", "National statistics/SUT", "economic/ OESA", "OESA guidance"],
            ["Ocean tourism", "Tourism statistics", "tourism/ OTSA", "OTSA guidance"],
            ["Waste/pollution", "Waste/water quality data", "waste/", "SEEA-CF templates"],
            ["Social/governance", "Census, surveys", "social/ scoping", "GOAP social framework"],
            ["Don't know", "Any", "Diagnostic tool", "diagnostic-tool.md"],
        ])

    add_heading_h2(doc, "Step 3: Walk through with them")
    add_body(doc, "1. Open the slides on screen (extent-slides / condition-slides / services-slides)")
    add_body(doc, "2. Follow the step-by-step guide for their pathway")
    add_body(doc, "3. Give them the exercise to practice")
    add_body(doc, "4. Open the Excel template and help them fill it with their own data")

    add_heading_h2(doc, "Facilitator specializations")
    make_goap_table(doc,
        ["Facilitator", "Specialization", "Availability"],
        [
            ["Jordan Gacutan", "Account structure, SEEA-EA, services, blue finance, GEF, NDC/NBSAP", "Day 2"],
            ["Maria Alarcon-Blazquez", "SEEA-EA technical guidance, valuation, ecosystem services, policy briefs", "Days 2-3"],
            ["Becca", "Social ocean accounts, community data, distributional analysis", "Days 2-3"],
        ])

    add_heading_h2(doc, "Key reminders")
    for r in [
        "Fellows must complete the GOAP Metadata Sheet BEFORE their session.",
        "If they have no data, start with Pathway A (global datasets / value transfer).",
        "A single-period account is valid. They do not need two time points.",
        "Emphasize: do not let perfect be the enemy of good. Tier 1 is a valid start.",
    ]:
        p = doc.add_paragraph()
        styled_run(p, "-- ", size=10, color=GREEN, bold=True)
        styled_run(p, r, size=10, color=BODY)

    path = outdir / "clinic_triage_card.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =========================================================================
# 3. EXTENT EXERCISE (print-ready A4)
# =========================================================================
def generate_extent_exercise_docx(outdir):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    add_title(doc, "Extent Account Exercise", size=16)
    add_subtitle(doc, "Ocean Accounting Clinic -- S-S Exchange 2026")

    add_heading_h3(doc, "Scenario")
    add_body(doc, "You are building an extent account for a coastal district with three marine ecosystem types. Satellite imagery gives you the following data.")

    add_heading_h3(doc, "Opening period (2018)")
    make_goap_table(doc,
        ["Ecosystem type", "Area (ha)"],
        [
            ["Coral reefs", "1,240"],
            ["Seagrass meadows", "860"],
            ["Mangroves", "315"],
            ["Sand/rubble/deep water", "7,585"],
            ["Total accounting area", "10,000"],
        ])

    add_heading_h3(doc, "Closing period (2023)")
    make_goap_table(doc,
        ["Ecosystem type", "Area (ha)"],
        [
            ["Coral reefs", "1,180"],
            ["Seagrass meadows", "825"],
            ["Mangroves", "340"],
            ["Sand/rubble/deep water", "7,655"],
            ["Total accounting area", "10,000"],
        ])

    add_heading_h3(doc, "Change matrix")
    make_goap_table(doc,
        ["Opening \\ Closing", "Coral", "Seagrass", "Mangrove", "Other"],
        [
            ["Coral", "1,140", "15", "0", "85"],
            ["Seagrass", "10", "790", "5", "55"],
            ["Mangrove", "0", "0", "300", "15"],
            ["Other", "30", "20", "35", "7,500"],
        ])

    add_heading_h2(doc, "Part 1: Complete the extent account table")
    add_body(doc, "Using the data above, calculate additions, reductions, net change, and closing extent for each ecosystem type.")
    make_goap_table(doc,
        ["", "Coral reefs", "Seagrass", "Mangroves", "Other", "Total"],
        [
            ["Opening extent (ha)", "", "", "", "", "10,000"],
            ["Total additions (ha)", "", "", "", "", ""],
            ["Total reductions (ha)", "", "", "", "", ""],
            ["Net change (ha)", "", "", "", "", ""],
            ["Closing extent (ha)", "", "", "", "", "10,000"],
        ])
    add_body(doc, "Hint: Additions for coral = sum of the Coral column excluding the diagonal. Reductions = sum of the Coral row excluding the diagonal.")

    add_heading_h2(doc, "Part 2: Calculate percentage changes")
    make_goap_table(doc,
        ["Ecosystem", "Opening (ha)", "Closing (ha)", "Net change (ha)", "Change (%)"],
        [
            ["Coral reefs", "1,240", "1,180", "", ""],
            ["Seagrass", "860", "825", "", ""],
            ["Mangroves", "315", "340", "", ""],
        ])
    add_body(doc, "Formula: Change (%) = (Net change / Opening extent) x 100")

    add_heading_h2(doc, "Part 3: Interpret")
    for q in [
        "1. Which ecosystem gained the most area? Where did the new area come from?",
        "2. Which ecosystem lost the most area? What did it convert to?",
        "3. 85 ha of coral became 'Other' (sand/rubble). What might explain this?",
        "4. Mangroves gained 35 ha from 'Other.' What does this likely represent?",
        "5. Is the total accounting area the same for both periods? Why must this be true?",
    ]:
        add_body(doc, q)

    path = outdir / "extent_exercise_print.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =========================================================================
# 4. CONDITION EXERCISE (print-ready A4)
# =========================================================================
def generate_condition_exercise_docx(outdir):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    add_title(doc, "Condition Account Exercise", size=16)
    add_subtitle(doc, "Ocean Accounting Clinic -- S-S Exchange 2026")

    add_heading_h3(doc, "Scenario")
    add_body(doc, "You have coral reef monitoring data from 4 sites in a marine reserve. Surveys were done in 2019 (opening) and 2024 (closing). Three indicators were measured.")

    add_heading_h3(doc, "Reference levels")
    make_goap_table(doc,
        ["Indicator", "Reference", "Direction"],
        [
            ["Live coral cover (%)", "50%", "Higher is better"],
            ["Fleshy macroalgae (%)", "60%", "Higher is worse"],
            ["Fish biomass (kg/ha)", "500", "Higher is better"],
        ])

    add_heading_h3(doc, "Opening data (2019)")
    make_goap_table(doc,
        ["Site", "Coral cover (%)", "Macroalgae (%)", "Fish biomass (kg/ha)"],
        [
            ["Site A", "28.5", "22.0", "180"],
            ["Site B", "35.2", "18.5", "245"],
            ["Site C", "18.7", "38.0", "120"],
            ["Site D", "42.1", "12.0", "310"],
        ])

    add_heading_h3(doc, "Closing data (2024)")
    make_goap_table(doc,
        ["Site", "Coral cover (%)", "Macroalgae (%)", "Fish biomass (kg/ha)"],
        [
            ["Site A", "22.0", "28.5", "160"],
            ["Site B", "30.8", "24.0", "220"],
            ["Site C", "15.2", "42.0", "105"],
            ["Site D", "38.5", "15.0", "290"],
        ])

    add_heading_h2(doc, "Part 1: Normalize to Condition Index")
    add_body(doc, "Higher-is-better: CI = Measured / Reference. Higher-is-worse: CI = 1 - (Measured / Reference). CI is capped at 0 to 1.")

    add_body(doc, "Opening (2019):")
    make_goap_table(doc,
        ["Site", "Coral CI", "Macroalgae CI", "Fish CI"],
        [
            ["Site A", "28.5 / 50 = ___", "1-(22/60) = ___", "180 / 500 = ___"],
            ["Site B", "", "", ""],
            ["Site C", "", "", ""],
            ["Site D", "", "", ""],
        ])

    add_body(doc, "Closing (2024):")
    make_goap_table(doc,
        ["Site", "Coral CI", "Macroalgae CI", "Fish CI"],
        [["Site A", "", "", ""], ["Site B", "", "", ""], ["Site C", "", "", ""], ["Site D", "", "", ""]])

    add_heading_h2(doc, "Part 2: Aggregate (mean of 4 sites)")
    make_goap_table(doc,
        ["", "Coral CI", "Macroalgae CI", "Fish CI"],
        [["Opening mean", "", "", ""], ["Closing mean", "", "", ""], ["Change", "", "", ""], ["Direction", "", "", ""]])

    add_heading_h2(doc, "Part 3: Condition account table")
    make_goap_table(doc,
        ["Indicator", "Reference", "Opening CI", "Closing CI", "Change", "Direction"],
        [
            ["Live coral cover", "50%", "", "", "", ""],
            ["Macroalgae cover", "60%", "", "", "", ""],
            ["Fish biomass", "500 kg/ha", "", "", "", ""],
        ])

    add_heading_h2(doc, "Part 4: Interpret")
    for q in [
        "1. Which indicator shows the most decline? What might drive it?",
        "2. Is the macroalgae CI improving or declining? What does this mean ecologically?",
        "3. Site C consistently has the lowest condition. What would you recommend?",
        "4. What is the composite condition index (average of 3 CIs) for opening and closing?",
    ]:
        add_body(doc, q)

    path = outdir / "condition_exercise_print.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =========================================================================
# 5. SERVICES EXERCISE (print-ready A4)
# =========================================================================
def generate_services_exercise_docx(outdir):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    add_title(doc, "Ecosystem Services Exercise", size=16)
    add_subtitle(doc, "Ocean Accounting Clinic -- S-S Exchange 2026")

    add_heading_h3(doc, "Scenario")
    add_body(doc, "You are building ecosystem service accounts for a coastal district with coral reefs (1,200 ha), mangroves (350 ha), and seagrass (800 ha).")

    add_heading_h3(doc, "Given data")
    make_goap_table(doc,
        ["Item", "Value"],
        [
            ["Reef fish catch", "180,000 kg/yr"],
            ["Average fish price", "USD 3.50/kg"],
            ["Total fishing costs", "USD 420,000/yr"],
            ["Mangrove NCP rate", "3.7 Mg CO2/ha/yr"],
            ["Seagrass NCP rate", "1.3 Mg CO2/ha/yr"],
            ["Social cost of carbon", "USD 51/Mg CO2"],
            ["Reef visitors", "15,000/yr"],
            ["Reef spending per visitor", "USD 85"],
            ["Protected coastline", "12,000 m"],
            ["Seawall replacement cost", "USD 5,000/m"],
            ["Seawall lifespan", "50 years"],
        ])

    add_heading_h2(doc, "Part 1: Fish provisioning")
    make_goap_table(doc,
        ["Step", "Calculation", "Value"],
        [
            ["Gross revenue", "180,000 x 3.50", "USD ___"],
            ["Total costs", "(given)", "USD 420,000"],
            ["Resource rent", "Revenue - Costs", "USD ___"],
        ])

    add_heading_h2(doc, "Part 2: Carbon sequestration")
    make_goap_table(doc,
        ["", "Mangroves", "Seagrass"],
        [
            ["Extent (ha)", "350", "800"],
            ["NCP rate (Mg CO2/ha/yr)", "3.7", "1.3"],
            ["Physical supply (Mg CO2/yr)", "", ""],
            ["Value at USD 51 (USD/yr)", "", ""],
        ])
    add_body(doc, "Total carbon value: USD ___/yr")

    add_heading_h2(doc, "Part 3: Recreation")
    make_goap_table(doc,
        ["Step", "Calculation", "Value"],
        [["Recreation value", "15,000 x USD 85", "USD ___/yr"]])

    add_heading_h2(doc, "Part 4: Coastal protection")
    make_goap_table(doc,
        ["Step", "Calculation", "Value"],
        [
            ["Total replacement cost", "12,000 m x USD 5,000", "USD ___"],
            ["Annualized (50 yr)", "Total / 50", "USD ___/yr"],
        ])

    add_heading_h2(doc, "Part 5: Summary table")
    make_goap_table(doc,
        ["Service", "Method", "Value (USD/yr)"],
        [
            ["Fish provisioning", "Resource rent", ""],
            ["Carbon sequestration", "Social cost of carbon", ""],
            ["Recreation", "Direct expenditure", ""],
            ["Coastal protection", "Replacement cost", ""],
            ["TOTAL", "", ""],
        ])

    add_heading_h2(doc, "Part 6: Interpret")
    for q in [
        "1. Which service has the highest value? Does this match your intuition?",
        "2. If carbon price were USD 15 instead of USD 51, what would the carbon value be?",
        "3. Does a seawall truly replace all services a coral reef provides? Why or why not?",
        "4. If fishing costs rose to USD 600,000, what happens to the resource rent?",
    ]:
        add_body(doc, q)

    path = outdir / "services_exercise_print.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =========================================================================
# MAIN
# =========================================================================
def main():
    outdir = Path(__file__).parent
    print("Generating print-ready GOAP Word documents...")

    generate_station_signs(outdir)
    generate_triage_card(outdir)
    generate_extent_exercise_docx(outdir)
    generate_condition_exercise_docx(outdir)
    generate_services_exercise_docx(outdir)

    print("\nDone! All print-ready files in:", outdir)


if __name__ == "__main__":
    main()
