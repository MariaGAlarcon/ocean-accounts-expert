#!/usr/bin/env python3
"""
Format QA script for GOAP print-notebook .docx files.

Checks and fixes:
1. Markdown artifacts: literal "**" in text -> convert to bold formatting
2. Orphaned titles: add keep_with_next before tables
3. Table column widths: enforce minimum widths for wide tables (>6 cols)
4. Empty paragraphs: collapse consecutive blank paragraphs
5. Consistent fonts: body text = Arial 10pt, headings per GOAP spec
"""

import os
import sys
import re
import glob
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# GOAP heading font sizes
# ---------------------------------------------------------------------------
HEADING_SIZES = {
    'Heading 1': Pt(16),
    'Heading 2': Pt(14),
    'Heading 3': Pt(12),
    'Heading 4': Pt(11),
}
BODY_FONT = 'Arial'
BODY_SIZE = Pt(10)
MIN_COL_WIDTH = Cm(1.5)  # minimum column width for wide tables


class QAReport:
    """Accumulates issues found and fixed for a single file."""
    def __init__(self, filepath):
        self.filepath = filepath
        self.issues = []

    def log(self, category, message):
        self.issues.append((category, message))

    def summary(self):
        if not self.issues:
            return f"  [CLEAN] {os.path.basename(self.filepath)}: no issues found."
        lines = [f"  {os.path.basename(self.filepath)}: {len(self.issues)} issue(s)"]
        for cat, msg in self.issues:
            lines.append(f"    [{cat}] {msg}")
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# 1. Fix markdown bold artifacts  (**text**)
# ---------------------------------------------------------------------------
def fix_markdown_bold(doc, report):
    """Find literal '**' in paragraph runs; remove them and apply bold."""
    count = 0
    for para in doc.paragraphs:
        full_text = para.text
        if '**' not in full_text:
            continue

        # Rebuild runs: collect all run texts, find ** pairs, set bold
        # Strategy: concatenate run texts, locate ** pairs in concatenated
        # string, then map character positions back to runs.
        run_texts = [r.text for r in para.runs]
        concat = ''.join(run_texts)

        if '**' not in concat:
            continue

        # Find all **...** spans
        pattern = re.compile(r'\*\*(.+?)\*\*')
        matches = list(pattern.finditer(concat))
        if not matches:
            # Orphaned ** without pairs — still strip them
            if '**' in concat:
                for run in para.runs:
                    if '**' in run.text:
                        run.text = run.text.replace('**', '')
                        count += 1
                        report.log("MARKDOWN", f"Removed orphaned '**' in: {full_text[:60]}...")
            continue

        # Build a bold-map: for each char position, should it be bold?
        bold_map = [False] * len(concat)
        strip_positions = set()
        for m in matches:
            # Mark the inner text as bold
            for i in range(m.start(1), m.end(1)):
                bold_map[i] = True
            # Mark the ** delimiters for stripping
            strip_positions.update(range(m.start(), m.start() + 2))
            strip_positions.update(range(m.end() - 2, m.end()))

        # Now rebuild runs with correct bold and stripped **
        # Map each char back to its run
        char_run_map = []
        for ri, rt in enumerate(run_texts):
            for ci in range(len(rt)):
                char_run_map.append(ri)

        # Group consecutive chars with same (run_index, bold) into new runs
        # But simpler: clear existing runs and create new ones
        # Preserve formatting of the first run as template
        if not para.runs:
            continue

        template_run = para.runs[0]
        # Collect original run formatting per run index
        run_formats = []
        for r in para.runs:
            run_formats.append({
                'bold': r.bold,
                'italic': r.italic,
                'underline': r.underline,
                'font_name': r.font.name,
                'font_size': r.font.size,
                'font_color_rgb': r.font.color.rgb if r.font.color and r.font.color.rgb else None,
            })

        # Build new segments: (text, bold_override, original_run_index)
        segments = []
        for pos in range(len(concat)):
            if pos in strip_positions:
                continue
            ri = char_run_map[pos]
            need_bold = bold_map[pos]
            ch = concat[pos]
            if segments and segments[-1][2] == ri and segments[-1][1] == need_bold:
                segments[-1] = (segments[-1][0] + ch, need_bold, ri)
            else:
                segments.append((ch, need_bold, ri))

        # Remove all existing runs
        for r in para.runs:
            r._element.getparent().remove(r._element)

        # Add new runs
        for text, make_bold, ri in segments:
            new_run = para.add_run(text)
            fmt = run_formats[ri] if ri < len(run_formats) else run_formats[0]
            if make_bold:
                new_run.bold = True
            else:
                new_run.bold = fmt['bold']
            new_run.italic = fmt['italic']
            new_run.underline = fmt['underline']
            if fmt['font_name']:
                new_run.font.name = fmt['font_name']
            if fmt['font_size']:
                new_run.font.size = fmt['font_size']
            if fmt['font_color_rgb']:
                new_run.font.color.rgb = fmt['font_color_rgb']

        count += len(matches)
        report.log("MARKDOWN", f"Fixed {len(matches)} bold marker(s) in: {full_text[:60]}...")

    # Also check table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    if '**' not in full_text:
                        continue
                    for run in para.runs:
                        if '**' in run.text:
                            # Simple approach for table cells: strip ** and bold the run
                            if re.search(r'\*\*.+?\*\*', run.text):
                                run.text = run.text.replace('**', '')
                                run.bold = True
                                count += 1
                                report.log("MARKDOWN", f"Fixed bold in table cell: {full_text[:50]}...")
                            else:
                                run.text = run.text.replace('**', '')
                                count += 1
                                report.log("MARKDOWN", f"Removed orphaned '**' in table cell: {full_text[:50]}...")

    return count


# ---------------------------------------------------------------------------
# 2. Keep titles with following tables (orphaned titles)
# ---------------------------------------------------------------------------
def fix_orphaned_titles(doc, report):
    """Add keep_with_next to paragraphs immediately before a table."""
    count = 0
    body = doc.element.body
    children = list(body)

    for i, child in enumerate(children):
        # Check if this is a table element
        if child.tag == qn('w:tbl'):
            # Look at the element immediately before
            if i > 0:
                prev = children[i - 1]
                if prev.tag == qn('w:p'):
                    # Check if it has text (potential title)
                    texts = prev.findall('.//' + qn('w:t'))
                    text_content = ''.join(t.text or '' for t in texts).strip()
                    if text_content and len(text_content) < 200:
                        # Add keep_with_next
                        pPr = prev.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                            prev.insert(0, pPr)
                        kwn = pPr.find(qn('w:keepNext'))
                        if kwn is None:
                            pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
                            count += 1
                            report.log("ORPHAN-TITLE", f"Added keep_with_next before table: '{text_content[:50]}...'")

            # Also check 2 elements back (title + blank para + table)
            if i > 1:
                prev2 = children[i - 2]
                prev1 = children[i - 1]
                if (prev2.tag == qn('w:p') and prev1.tag == qn('w:p')):
                    # Check if prev1 is empty and prev2 has short text
                    texts1 = prev1.findall('.//' + qn('w:t'))
                    text1 = ''.join(t.text or '' for t in texts1).strip()
                    texts2 = prev2.findall('.//' + qn('w:t'))
                    text2 = ''.join(t.text or '' for t in texts2).strip()
                    if not text1 and text2 and len(text2) < 200:
                        pPr = prev2.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                            prev2.insert(0, pPr)
                        kwn = pPr.find(qn('w:keepNext'))
                        if kwn is None:
                            pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
                            count += 1
                            report.log("ORPHAN-TITLE", f"Added keep_with_next (with gap): '{text2[:50]}...'")

    return count


# ---------------------------------------------------------------------------
# 3. Table column widths — enforce minimum for wide tables
# ---------------------------------------------------------------------------
def fix_table_widths(doc, report):
    """For tables with >6 columns, ensure no column is unreasonably narrow."""
    count = 0
    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        ncols = len(table.columns)
        if ncols <= 6:
            continue

        for col_idx, col in enumerate(table.columns):
            # Check width
            current_width = col.width
            # Use a small tolerance (1000 EMU ~ 0.03mm) to avoid oscillation
            if current_width is not None and current_width < (MIN_COL_WIDTH - 1000):
                col.width = MIN_COL_WIDTH
                count += 1
                report.log("COL-WIDTH",
                           f"Table {ti+1}: col {col_idx+1} widened from "
                           f"{current_width} to {MIN_COL_WIDTH}")

            # Also enforce via cell widths in first row
            if table.rows:
                try:
                    cell = table.rows[0].cells[col_idx]
                    tc = cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        tcW = tcPr.find(qn('w:tcW'))
                        if tcW is not None:
                            w_val = tcW.get(qn('w:w'))
                            w_type = tcW.get(qn('w:type'))
                            if w_type == 'dxa' and w_val:
                                width_emu = int(w_val) * 635  # twips to EMU approx
                                if width_emu < (MIN_COL_WIDTH - 1000):
                                    new_twips = int(MIN_COL_WIDTH / 635)
                                    tcW.set(qn('w:w'), str(new_twips))
                                    count += 1
                                    report.log("COL-WIDTH",
                                               f"Table {ti+1}: cell col {col_idx+1} width "
                                               f"adjusted in XML")
                except (IndexError, AttributeError):
                    pass

    return count


# ---------------------------------------------------------------------------
# 4. Remove consecutive empty paragraphs
# ---------------------------------------------------------------------------
def fix_empty_paragraphs(doc, report):
    """Remove consecutive empty paragraphs (keep at most 1)."""
    count = 0
    body = doc.element.body
    children = list(body)

    consecutive_empties = []
    for child in children:
        if child.tag == qn('w:p'):
            texts = child.findall('.//' + qn('w:t'))
            text_content = ''.join(t.text or '' for t in texts).strip()
            # Also check for images/drawings — don't remove those
            has_drawing = child.findall('.//' + qn('w:drawing'))
            has_pict = child.findall('.//' + qn('w:pict'))

            if not text_content and not has_drawing and not has_pict:
                consecutive_empties.append(child)
            else:
                # Process accumulated empties
                if len(consecutive_empties) > 1:
                    # Remove all but the first
                    for empty in consecutive_empties[1:]:
                        body.remove(empty)
                        count += 1
                consecutive_empties = []
        else:
            if len(consecutive_empties) > 1:
                for empty in consecutive_empties[1:]:
                    body.remove(empty)
                    count += 1
            consecutive_empties = []

    # Handle trailing empties
    if len(consecutive_empties) > 1:
        for empty in consecutive_empties[1:]:
            body.remove(empty)
            count += 1

    if count:
        report.log("EMPTY-PARA", f"Removed {count} consecutive empty paragraph(s)")
    return count


# ---------------------------------------------------------------------------
# 5. Consistent fonts
# ---------------------------------------------------------------------------
def fix_fonts(doc, report):
    """Ensure body text is Arial 10pt; headings follow GOAP spec."""
    count = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''

        # Handle headings
        if style_name in HEADING_SIZES:
            target_size = HEADING_SIZES[style_name]
            for run in para.runs:
                if run.font.name != BODY_FONT:
                    run.font.name = BODY_FONT
                    # Also set East Asian and Complex Script fonts
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            rFonts.set(qn('w:ascii'), BODY_FONT)
                            rFonts.set(qn('w:hAnsi'), BODY_FONT)
                            rFonts.set(qn('w:cs'), BODY_FONT)
                    count += 1
                    report.log("FONT", f"Heading run font -> Arial: '{run.text[:40]}...'")
                if run.font.size != target_size:
                    old_size = run.font.size
                    run.font.size = target_size
                    count += 1
                    report.log("FONT", f"Heading '{style_name}' size {old_size} -> {target_size}: '{run.text[:30]}...'")

        # Handle body text (Normal, List Paragraph, Body Text, etc.)
        elif style_name in ('Normal', 'Body Text', 'List Paragraph',
                            'List Bullet', 'List Number', 'No Spacing',
                            'Body Text 2', 'Body Text 3'):
            for run in para.runs:
                if run.font.name and run.font.name != BODY_FONT:
                    old_name = run.font.name
                    run.font.name = BODY_FONT
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            rFonts.set(qn('w:ascii'), BODY_FONT)
                            rFonts.set(qn('w:hAnsi'), BODY_FONT)
                            rFonts.set(qn('w:cs'), BODY_FONT)
                    count += 1
                    report.log("FONT", f"Body font {old_name} -> Arial: '{run.text[:40]}...'")
                if run.font.size and run.font.size != BODY_SIZE:
                    old_size = run.font.size
                    run.font.size = BODY_SIZE
                    count += 1
                    report.log("FONT", f"Body size {old_size} -> {BODY_SIZE}: '{run.text[:40]}...'")

    # Also fix fonts in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name and run.font.name != BODY_FONT:
                            old_name = run.font.name
                            run.font.name = BODY_FONT
                            rPr = run._element.find(qn('w:rPr'))
                            if rPr is not None:
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is not None:
                                    rFonts.set(qn('w:ascii'), BODY_FONT)
                                    rFonts.set(qn('w:hAnsi'), BODY_FONT)
                                    rFonts.set(qn('w:cs'), BODY_FONT)
                            count += 1

    return count


# ---------------------------------------------------------------------------
# Main processing
# ---------------------------------------------------------------------------
def process_file(filepath):
    """Run all QA checks on a single .docx file."""
    report = QAReport(filepath)
    try:
        doc = Document(filepath)
    except Exception as e:
        report.log("ERROR", f"Could not open file: {e}")
        return report, False

    total_fixes = 0
    total_fixes += fix_markdown_bold(doc, report)
    total_fixes += fix_orphaned_titles(doc, report)
    total_fixes += fix_table_widths(doc, report)
    total_fixes += fix_empty_paragraphs(doc, report)
    total_fixes += fix_fonts(doc, report)

    if total_fixes > 0:
        try:
            doc.save(filepath)
            report.log("SAVED", f"File saved with {total_fixes} fix(es)")
        except Exception as e:
            report.log("ERROR", f"Could not save: {e}")
            return report, False
    return report, total_fixes > 0


def process_folder(folder_path):
    """Process all .docx files in a folder."""
    docx_files = sorted(glob.glob(os.path.join(folder_path, '*.docx')))
    if not docx_files:
        print(f"  (no .docx files found)")
        return 0, 0

    total_files = 0
    total_fixed = 0
    for fp in docx_files:
        report, was_fixed = process_file(fp)
        print(report.summary())
        total_files += 1
        if was_fixed:
            total_fixed += 1
    return total_files, total_fixed


def main():
    base = os.path.dirname(os.path.abspath(__file__))

    # Determine which folders to process
    if len(sys.argv) > 1:
        folders = sys.argv[1:]
    else:
        # Process all numbered folders
        folders = sorted(glob.glob(os.path.join(base, '[0-9][0-9]-*')))

    grand_files = 0
    grand_fixed = 0

    for folder in folders:
        folder = os.path.abspath(folder)
        if not os.path.isdir(folder):
            print(f"Skipping (not a directory): {folder}")
            continue
        print(f"\n{'='*70}")
        print(f"FOLDER: {os.path.basename(folder)}")
        print(f"{'='*70}")
        nf, nfixed = process_folder(folder)
        grand_files += nf
        grand_fixed += nfixed

    print(f"\n{'='*70}")
    print(f"GRAND TOTAL: {grand_files} files checked, {grand_fixed} files modified")
    print(f"{'='*70}")


if __name__ == '__main__':
    main()
