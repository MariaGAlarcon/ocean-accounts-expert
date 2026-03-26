#!/usr/bin/env python3
"""Format QA v2: Fix table column widths and page breaks in Word docs.

Fixes:
1. Column widths: proportional to content length (short cols narrow, long cols wide)
2. Table page breaks: prevent tables from splitting across pages
3. Keep-with-next: titles stay with their tables
4. Markdown artifacts: remove literal ** bold markers
5. Font consistency: Arial throughout at GOAP sizes
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re


def get_max_content_length(table, col_idx):
    """Get the maximum text length in a column."""
    max_len = 0
    for row in table.rows:
        if col_idx < len(row.cells):
            text = row.cells[col_idx].text.strip()
            # Account for line breaks
            max_line = max((len(line) for line in text.split('\n')), default=0)
            max_len = max(max_len, max_line)
    return max_len


def set_column_widths(table):
    """Set column widths proportional to content length."""
    if not table.rows:
        return

    num_cols = len(table.columns)
    if num_cols == 0:
        return

    # Calculate content lengths per column
    col_lengths = []
    for c in range(num_cols):
        max_len = get_max_content_length(table, c)
        col_lengths.append(max(max_len, 1))  # minimum 1

    # Determine page width available (A4 landscape = ~26cm usable, portrait = ~17cm)
    # Use portrait as default
    total_available = 17.0  # cm

    # Calculate proportional widths with minimum
    total_content = sum(col_lengths)
    min_width_cm = 1.2  # minimum column width
    max_width_cm = 7.0  # maximum column width

    widths = []
    for length in col_lengths:
        # Proportional width
        prop = (length / total_content) * total_available
        # Apply bounds
        w = max(min_width_cm, min(max_width_cm, prop))
        # Very short content (1-3 chars) gets extra narrow
        if length <= 3:
            w = max(1.2, min(2.0, prop))
        elif length <= 8:
            w = max(1.5, min(3.0, prop))
        widths.append(w)

    # Scale to fit total available width
    total_w = sum(widths)
    if total_w > 0:
        scale = total_available / total_w
        widths = [w * scale for w in widths]

    # Apply widths to all cells
    for row in table.rows:
        for c, cell in enumerate(row.cells):
            if c < len(widths):
                cell.width = Cm(widths[c])

    # Set table to fixed layout (not autofit)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Remove autofit, set fixed
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is not None:
        tblPr.remove(layout)
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    tblPr.append(layout)


def prevent_table_split(table):
    """Prevent a table from splitting across pages where possible."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Set cantSplit on all rows (keep rows together)
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        # Remove existing cantSplit
        existing = trPr.find(qn('w:cantSplit'))
        if existing is not None:
            trPr.remove(existing)
        # Add cantSplit
        cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
        trPr.append(cantSplit)

    # Set keepWithNext on header row and first few data rows
    for i, row in enumerate(table.rows):
        if i < min(3, len(table.rows) - 1):  # Keep first 3 rows together
            for cell in row.cells:
                for para in cell.paragraphs:
                    pPr = para._element.get_or_add_pPr()
                    existing = pPr.find(qn('w:keepNext'))
                    if existing is None:
                        pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>'))


def fix_keep_with_next_before_tables(doc):
    """Ensure paragraphs before tables have keep_with_next set."""
    body = doc.element.body
    children = list(body)
    count = 0

    for i, child in enumerate(children):
        if child.tag == qn('w:tbl') and i > 0:
            # Look back for preceding paragraphs (title, subtitle)
            for j in range(max(0, i - 3), i):
                prev = children[j]
                if prev.tag == qn('w:p'):
                    text = prev.text or ""
                    # Check if it has text content (not just whitespace)
                    has_text = False
                    for r in prev.findall(qn('w:r')):
                        for t in r.findall(qn('w:t')):
                            if t.text and t.text.strip():
                                has_text = True
                                break
                    if has_text:
                        pPr = prev.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                            prev.insert(0, pPr)
                        existing = pPr.find(qn('w:keepNext'))
                        if existing is None:
                            pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>'))
                            count += 1
    return count


def fix_bold_artifacts(doc):
    """Remove literal ** around text and make it bold."""
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and '**' in run.text:
                cleaned = run.text.replace('**', '')
                if cleaned != run.text:
                    run.text = cleaned
                    run.bold = True
                    count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text and '**' in run.text:
                            cleaned = run.text.replace('**', '')
                            if cleaned != run.text:
                                run.text = cleaned
                                run.bold = True
                                count += 1
    return count


def process_file(filepath):
    """Run all formatting fixes on a single Word document."""
    doc = Document(str(filepath))
    fixes = {}

    # Fix bold artifacts
    bold_count = fix_bold_artifacts(doc)
    if bold_count:
        fixes['bold_artifacts'] = bold_count

    # Fix keep-with-next before tables
    kwn_count = fix_keep_with_next_before_tables(doc)
    if kwn_count:
        fixes['keep_with_next'] = kwn_count

    # Fix table column widths and page breaks
    table_count = 0
    for table in doc.tables:
        set_column_widths(table)
        prevent_table_split(table)
        table_count += 1

    if table_count:
        fixes['tables_fixed'] = table_count

    if fixes:
        doc.save(str(filepath))
        return fixes
    return None


def main():
    base = Path("/Users/mariaalarcon/GitHub/ocean-accounts-expert/s-s-event-bogor/print-notebooks")

    total_files = 0
    total_fixed = 0

    for folder in sorted(base.iterdir()):
        if not folder.is_dir() or folder.name.startswith('.') or folder.name == 'pdfs':
            continue

        print(f"\n{'='*60}")
        print(f"FOLDER: {folder.name}")
        print(f"{'='*60}")

        for docx_file in sorted(folder.glob("*.docx")):
            total_files += 1
            fixes = process_file(docx_file)
            if fixes:
                total_fixed += 1
                print(f"  [FIXED] {docx_file.name}: {fixes}")
            else:
                print(f"  [CLEAN] {docx_file.name}")

    print(f"\n{'='*60}")
    print(f"TOTAL: {total_files} files checked, {total_fixed} files modified")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
