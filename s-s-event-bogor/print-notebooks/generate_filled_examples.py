#!/usr/bin/env python3
"""Generate GOAP-styled Word documents with DETAILED filled-in example accounting tables.

One document per account type, each with realistic data including IUCN GET codes,
species names, ISIC codes, actual reference levels, and real country data.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── GOAP colour palette ──────────────────────────────────────────────
TEAL = RGBColor(0x0A, 0x54, 0x55)
GREEN = RGBColor(0x3B, 0x9C, 0x7B)
BODY = RGBColor(0x30, 0x30, 0x2F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_HEX = "404040"

# ── Shared helpers ────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:left w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:right w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def styled_run(p, text, size=10, bold=False, color=BODY):
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def add_title(doc, text):
    p = doc.add_paragraph()
    styled_run(p, text, size=16, color=TEAL)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    styled_run(p, text, size=11, color=GREEN)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    styled_run(p, text, size=9, color=RGBColor(0x55, 0x55, 0x55))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_interpretation(doc, heading, text):
    """Add a bold heading followed by an interpretation paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    styled_run(p, heading, size=10, bold=True, color=TEAL)
    p2 = doc.add_paragraph()
    styled_run(p2, text, size=10, color=BODY)
    p2.paragraph_format.space_after = Pt(6)
    return p2


def make_table(doc, headers, rows, row_header_col=0):
    """Create a GOAP-styled table with green headers and teal row headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, h, size=10, bold=True, color=WHITE)
        set_cell_shading(cell, "3B9C7B")

    # Data rows
    for r, row_data in enumerate(rows, 1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]

            if c == row_header_col:
                styled_run(p, str(val), size=10, bold=True, color=WHITE)
                set_cell_shading(cell, "0A5455")
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT
                styled_run(p, str(val) if val else "--", size=10, color=BODY)

            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    set_table_borders(table)
    return table


def set_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_w, new_h = section.page_height, section.page_width
    section.page_width = new_w
    section.page_height = new_h
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


# =====================================================================
# 1. EXTENT -- Laamu Atoll, Maldives
# =====================================================================
def generate_extent(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Worked Example: Extent Account -- Laamu Atoll, Maldives")
    add_subtitle(doc, "Accounting area: Laamu Atoll  |  Opening year: 2017  |  Closing year: 2020")
    add_note(doc, "Data sources: Allen Coral Atlas (2020), Maldives Marine Research Institute habitat maps, "
             "Blue Prosperity Coalition surveys. IUCN Global Ecosystem Typology (GET) codes used throughout.")

    # ── Table 1: Extent Account ──
    doc.add_paragraph()
    add_title(doc, "Table 1: Ecosystem Extent Account (hectares)")
    make_table(doc,
        ["", "Photic coral\nreefs (M1.3)", "Seagrass\nmeadows (M1.1)", "Mangroves\n(MFT1.2)",
         "Sand / rubble /\ndeep water", "Total"],
        [
            ["Opening extent (2017)", "7,420", "4,856", "19", "149,653", "161,948"],
            ["Additions: natural expansion", "0", "32", "0", "0", "32"],
            ["Additions: managed restoration", "0", "0", "1", "0", "1"],
            ["Additions: reclassification", "0", "0", "0", "0", "0"],
            ["Total additions", "0", "32", "1", "0", "33"],
            ["Reductions: natural reduction", "-18", "-12", "0", "0", "-30"],
            ["Reductions: managed conversion", "-7", "-20", "-1", "0", "-28"],
            ["Reductions: reclassification", "0", "0", "0", "0", "0"],
            ["Total reductions", "-25", "-32", "-1", "0", "-58"],
            ["Net change", "-25", "0", "0", "+25", "0"],
            ["Closing extent (2020)", "7,395", "4,856", "19", "149,678", "161,948"],
        ])

    add_interpretation(doc, "Interpretation:",
        "The total accounting area (161,948 ha) remains constant between opening and closing -- "
        "this is an accounting identity. Coral reefs declined by 25 ha (0.3%), primarily from "
        "natural reduction linked to the 2016 mass-bleaching event's lagged mortality. Seagrass "
        "showed no net change: natural expansion (32 ha) was exactly offset by reductions. "
        "Mangroves are a very small component (19 ha) and remained stable. The 25 ha lost from "
        "reef and seagrass was reclassified to sand/rubble/deep water.")

    # ── Table 2: Change Matrix ──
    doc.add_page_break()
    add_title(doc, "Table 2: Ecosystem Extent Change Matrix (hectares, 2017-2020)")
    add_note(doc, "Diagonal cells (bold) = area that stayed the same type. Off-diagonal = transitions.")

    make_table(doc,
        ["Opening \\ Closing", "Photic coral\nreefs (M1.3)", "Seagrass\nmeadows (M1.1)",
         "Mangroves\n(MFT1.2)", "Sand / rubble /\ndeep water", "Opening\ntotal"],
        [
            ["Photic coral reefs (M1.3)", "7,395", "0", "0", "25", "7,420"],
            ["Seagrass meadows (M1.1)", "0", "4,824", "0", "32", "4,856"],
            ["Mangroves (MFT1.2)", "0", "0", "18", "1", "19"],
            ["Sand / rubble / deep water", "0", "32", "1", "149,620", "149,653"],
            ["Closing total", "7,395", "4,856", "19", "149,678", "161,948"],
        ])

    add_interpretation(doc, "Reading the change matrix:",
        "Read across rows to see where opening area went. For example, of the 7,420 ha of coral reef "
        "in 2017, 7,395 ha stayed as reef and 25 ha converted to sand/rubble. Read down columns to "
        "see where closing area came from. The diagonal shows stable area; off-diagonal shows "
        "transitions. Note that seagrass gained 32 ha from sand but lost 32 ha back -- so net change "
        "was zero, even though 64 ha of transitions occurred.")

    path = outdir / "01-extent" / "10-filled-example.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# =====================================================================
# 2. CONDITION -- Belize Coral Reefs
# =====================================================================
def generate_condition(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Worked Example: Condition Account -- Belize Coral Reefs")
    add_subtitle(doc, "Ecosystem type: Photic coral reefs (IUCN GET M1.3)  |  "
                 "Accounting area: Northern Barrier Complex, Belize")
    add_note(doc, "Data source: Atlantic and Gulf Rapid Reef Assessment (AGRRA) surveys, "
             "2019 (opening) and 2024 (closing). All indicators measured at permanent monitoring sites.")

    # ── Table 3: Condition Account ──
    doc.add_paragraph()
    add_title(doc, "Table 3: Ecosystem Condition Account for Photic Coral Reefs (M1.3)")
    add_note(doc, "CI = Condition Index. Higher-is-better: CI = Value / Reference. "
             "Higher-is-worse: CI = 1 - (Value / Reference). All CIs capped at [0, 1].")

    make_table(doc,
        ["Indicator", "Unit", "Direction", "Reference level\n(source)", "Opening\nvalue",
         "Opening\nCI", "Closing\nvalue", "Closing\nCI", "Change\nin CI", "Status"],
        [
            ["Live coral cover", "%", "Higher is\nbetter",
             "50%\n(Perry et al. 2013,\nIndo-Pacific)", "31.1", "0.62", "26.6", "0.53", "-0.09", "Declining"],
            ["Fleshy macroalgae\ncover", "%", "Higher is\nworse",
             "60%\n(maximum\ndegraded state)", "22.6", "0.62", "27.4", "0.54", "-0.08", "Declining"],
            ["Fish biomass", "kg/ha", "Higher is\nbetter",
             "500 kg/ha\n(McClanahan et al.\n2011, unfished)", "213.8", "0.43", "193.8", "0.39", "-0.04", "Declining"],
            ["Coral recruit\ndensity", "per m\u00b2", "Higher is\nbetter",
             "15 per m\u00b2\n(regional\nbenchmark)", "8.5", "0.57", "6.2", "0.41", "-0.15", "Declining"],
            ["Reef structural\nrelief", "cm", "Higher is\nbetter",
             "150 cm\n(healthy reef\nstructure)", "82.5", "0.55", "78.0", "0.52", "-0.03", "Stable"],
            ["Bleaching\nprevalence", "% colonies\nbleached", "Higher is\nworse",
             "100%\n(all colonies\nbleached)", "8.5", "0.92", "18.3", "0.82", "-0.10", "Declining"],
            ["Composite CI", "(average of\nabove)", "", "", "", "0.62", "", "0.53", "-0.08", "Declining"],
        ])

    add_interpretation(doc, "Data source:",
        "AGRRA surveys 2019 and 2024, Northern Barrier Complex, Belize. Surveys covered "
        "32 permanent transects across 8 reef sites spanning fore-reef and back-reef habitats.")

    add_interpretation(doc, "Interpretation:",
        "The composite Condition Index declined from 0.62 to 0.53 over five years, indicating "
        "a reef system moving further from reference condition. Coral recruit density showed the "
        "largest decline (-0.15), suggesting reduced coral reproduction and early survivorship -- "
        "a leading indicator of future reef decline. Bleaching prevalence more than doubled (8.5% "
        "to 18.3%), consistent with the 2023 marine heatwave that affected the Mesoamerican "
        "Barrier Reef. Fish biomass remained low (< 50% of reference) in both periods, reflecting "
        "chronic overfishing. Reef structural relief was the most stable indicator (-0.03), as "
        "physical structure changes slowly.")

    add_interpretation(doc, "How to read the CI:",
        "A CI of 1.0 means the indicator is at its reference level (best condition). A CI of 0.0 "
        "means fully degraded. For 'higher-is-worse' indicators (macroalgae, bleaching), the CI "
        "formula is inverted: CI = 1 - (value/reference), so higher observed values produce lower CIs.")

    path = outdir / "02-condition" / "11-filled-example.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# =====================================================================
# 3. SERVICES -- Coastal District
# =====================================================================
def generate_services(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Worked Example: Ecosystem Services Account -- Coastal District")
    add_subtitle(doc, "Accounting area: illustrative coastal district with coral reefs, "
                 "seagrass beds, and mangroves  |  Year: 2023")
    add_note(doc, "Three ecosystem types supply services to the local economy. "
             "Part A (physical units) is the core account; Part B (monetary) supplements it.")

    # ── Table 4a: Physical Supply ──
    doc.add_paragraph()
    add_title(doc, "Table 4a: Physical Supply of Ecosystem Services (Part A)")

    make_table(doc,
        ["Service", "SEEA EA\nreference list", "Unit", "Coral reefs\n(M1.3)", "Seagrass\n(M1.1)",
         "Mangroves\n(MFT1.2)", "Total"],
        [
            ["Fish provisioning", "Provisioning", "kg/yr", "120,000", "45,000", "15,000", "180,000"],
            ["Wood (fuel)", "Provisioning", "tonnes/yr", "--", "--", "85", "85"],
            ["Carbon sequestration", "Regulating", "Mg CO\u2082/yr", "--", "1,040", "1,295", "2,335"],
            ["Coastal protection", "Regulating", "m coastline\nprotected", "12,000", "--", "3,500", "15,500"],
            ["Nursery habitat", "Regulating", "kg biomass\nexported/yr", "4,500", "2,800", "--", "7,300"],
            ["Sediment retention /\ncarbonate production", "Regulating", "m\u00b3 CaCO\u2083/yr", "8,400", "--", "--", "8,400"],
            ["Recreation (reef\ndiving/snorkelling)", "Cultural", "visitors/yr", "15,000", "--", "--", "15,000"],
            ["Recreation (mangrove\nkayak/birdwatch)", "Cultural", "trips/yr", "--", "--", "2,400", "2,400"],
            ["Gleaning\n(invertebrate harvest)", "Cultural", "hours/yr", "--", "18,000", "--", "18,000"],
        ])

    add_note(doc, "Dashes (--) indicate that the ecosystem type does not supply the service "
             "or the flow is negligible. Carbon sequestration counts only annual burial, "
             "not standing stock.")

    # ── Table 4b: Monetary Supply ──
    doc.add_page_break()
    add_title(doc, "Table 4b: Monetary Supply of Ecosystem Services (Part B -- USD/year)")

    make_table(doc,
        ["Service", "Valuation\nmethod", "Value\ntype", "Coral reefs\n(M1.3)", "Seagrass\n(M1.1)",
         "Mangroves\n(MFT1.2)", "Total\n(USD/yr)"],
        [
            ["Fish provisioning", "Resource rent", "Market", "96,000", "36,000", "12,000", "144,000"],
            ["Wood (fuel)", "Market price", "Market", "--", "--", "5,100", "5,100"],
            ["Carbon sequestration", "Social cost of\ncarbon (USD 51/t)", "Non-market",
             "--", "53,040", "66,045", "119,085"],
            ["Coastal protection", "Avoided damage\ncost", "Non-market",
             "1,440,000", "--", "420,000", "1,860,000"],
            ["Nursery habitat", "Productivity\nchange method", "Market", "90,000", "56,000", "--", "146,000"],
            ["Sediment retention", "Replacement\ncost", "Non-market", "252,000", "--", "--", "252,000"],
            ["Recreation (reef)", "Direct\nexpenditure", "Market", "1,050,000", "--", "--", "1,050,000"],
            ["Recreation (mangrove)", "Direct\nexpenditure", "Market", "--", "--", "62,400", "62,400"],
            ["Gleaning", "Equivalent\nwage (USD 2.5/hr)", "Mixed", "--", "48,000", "--", "48,000"],
            ["TOTAL", "", "", "2,928,000", "193,040", "565,545", "3,686,585"],
        ])

    add_interpretation(doc, "Interpretation:",
        "Total annual value of ecosystem services is estimated at USD 3,686,585. Coastal protection "
        "(USD 1,860,000) and recreation (USD 1,112,400 combined) dominate total value, "
        "accounting for approximately 81% of the total. Fish provisioning -- often assumed to be "
        "the dominant service -- represents only about 4% of the monetary total (USD 144,000). "
        "This highlights a common finding: regulating and cultural services frequently exceed "
        "provisioning services in economic value, supporting the case for ecosystem conservation.")

    add_interpretation(doc, "Methodological notes:",
        "Fish provisioning valued at resource rent (market price minus production costs), not gross "
        "revenue. Carbon sequestration uses the US EPA social cost of carbon (USD 51/tCO\u2082, 2020 "
        "value). Coastal protection based on avoided damage to coastal infrastructure using storm "
        "damage functions. Gleaning valued at local equivalent wage rate as subsistence activity "
        "has no market price.")

    # ── Table 4c: Physical Use ──
    doc.add_page_break()
    add_title(doc, "Table 4c: Physical Use of Ecosystem Services")
    add_note(doc, "The use table shows which economic sectors benefit from each service. "
             "Total use must equal total supply from Table 4a for every service row.")

    make_table(doc,
        ["Service", "Unit", "Fisheries", "Tourism", "Coastal HH", "Govt", "Global", "Total"],
        [
            ["Fish provisioning", "kg/yr", "180,000", "0", "0", "0", "0", "180,000"],
            ["Carbon sequestration", "Mg CO\u2082/yr", "0", "0", "0", "0", "2,335", "2,335"],
            ["Coastal protection", "m", "0", "0", "10,000", "5,500", "0", "15,500"],
            ["Nursery habitat", "kg/yr", "7,300", "0", "0", "0", "0", "7,300"],
            ["Recreation", "visitors/yr", "0", "15,000", "2,400", "0", "0", "17,400"],
            ["Gleaning", "hours/yr", "0", "0", "18,000", "0", "0", "18,000"],
        ])

    add_interpretation(doc, "Interpretation:",
        "Fish provisioning is used entirely by the fisheries sector. Carbon sequestration "
        "benefits the global community through climate regulation. Coastal protection is "
        "split between coastal households (property protection, 10,000 m) and government "
        "(public infrastructure, 5,500 m). Recreation is used by tourism businesses "
        "(15,000 visitors) and coastal households for local recreation (2,400 visitors). "
        "Gleaning is a subsistence activity benefiting coastal households exclusively.")

    # ── Table 4d: Monetary Use ──
    doc.add_page_break()
    add_title(doc, "Table 4d: Monetary Use of Ecosystem Services (USD/year)")

    make_table(doc,
        ["Service", "Method", "Fisheries", "Tourism", "Coastal HH", "Govt", "Global", "Total"],
        [
            ["Fish provisioning", "Resource rent", "210,000", "0", "0", "0", "0", "210,000"],
            ["Carbon sequestration", "SCC", "0", "0", "0", "0", "119,085", "119,085"],
            ["Coastal protection", "Replacement", "0", "0", "880,000", "495,000", "0", "1,375,000"],
            ["Nursery habitat", "Productivity", "127,750", "0", "0", "0", "0", "127,750"],
            ["Recreation", "Expenditure", "0", "1,275,000", "84,000", "0", "0", "1,359,000"],
            ["Gleaning", "Equiv. wage", "0", "0", "63,000", "0", "0", "63,000"],
            ["TOTAL", "", "337,750", "1,275,000", "1,027,000", "495,000", "119,085", "3,253,835"],
        ])

    add_interpretation(doc, "Interpretation:",
        "Coastal households receive the most value (USD 1,027,000, 31.6%), driven by coastal "
        "protection and gleaning. Tourism is the second-largest beneficiary (USD 1,275,000, "
        "39.2%), almost entirely from recreation. Fisheries receive USD 337,750 from fish "
        "provisioning and nursery habitat combined. Government benefits from coastal protection "
        "of public infrastructure (USD 495,000). The global community receives USD 119,085 "
        "from carbon sequestration valued at the social cost of carbon.")

    # ── Table 4e: Integrated Supply-Use Table (SEEA EA Table 7.1 format) ──
    doc.add_page_break()
    add_title(doc, "Table 4e: Integrated Supply-Use Table (SEEA EA Table 7.1 format)")
    add_note(doc, "The integrated SUT shows the complete picture. Supply (top) shows ecosystems "
             "providing services. Use (bottom) shows who benefits. Total supply = Total use for "
             "each service.")
    add_note(doc, "The SEEA EA presents supply and use together in a single integrated table. "
             "The left half shows which ecosystem types supply each service; the right half shows "
             "which economic sectors use (benefit from) each service. Total supply = Total use for every row.")

    make_table(doc,
        ["Service", "Unit",
         "Coral reefs", "Seagrass", "Mangroves",
         "Fisheries", "Tourism", "Coastal HH", "Govt", "Global",
         "TOTAL"],
        [
            ["── SUPPLY ──", "", "", "", "", "", "", "", "", "", ""],
            ["Fish provisioning", "kg/yr", "120,000", "45,000", "15,000",
             "", "", "", "", "", "180,000"],
            ["Carbon sequestration", "Mg CO\u2082/yr", "0", "1,040", "1,295",
             "", "", "", "", "", "2,335"],
            ["Coastal protection", "m coastline", "12,000", "0", "3,500",
             "", "", "", "", "", "15,500"],
            ["Nursery habitat", "kg biomass/yr", "4,500", "2,800", "0",
             "", "", "", "", "", "7,300"],
            ["Recreation", "visitors/yr", "15,000", "0", "2,400",
             "", "", "", "", "", "17,400"],
            ["Gleaning", "hours/yr", "0", "18,000", "0",
             "", "", "", "", "", "18,000"],
            ["── USE ──", "", "", "", "", "", "", "", "", "", ""],
            ["Fish provisioning", "kg/yr", "", "", "",
             "180,000", "0", "0", "0", "0", "180,000"],
            ["Carbon sequestration", "Mg CO\u2082/yr", "", "", "",
             "0", "0", "0", "0", "2,335", "2,335"],
            ["Coastal protection", "m coastline", "", "", "",
             "0", "0", "10,000", "5,500", "0", "15,500"],
            ["Nursery habitat", "kg biomass/yr", "", "", "",
             "7,300", "0", "0", "0", "0", "7,300"],
            ["Recreation", "visitors/yr", "", "", "",
             "0", "15,000", "2,400", "0", "0", "17,400"],
            ["Gleaning", "hours/yr", "", "", "",
             "0", "0", "18,000", "0", "0", "18,000"],
        ])

    add_interpretation(doc, "Reading the integrated SUT:",
        "The integrated SUT shows the complete picture. Supply (top) shows ecosystems providing "
        "services. Use (bottom) shows who benefits. Total supply = Total use for each service. "
        "For example, fish provisioning: 120,000 kg from coral reefs + 45,000 from seagrass + "
        "15,000 from mangroves = 180,000 kg total supply. All 180,000 kg is used by the "
        "fisheries sector = 180,000 kg total use. This accounting identity must hold for "
        "every row and is a fundamental check on the account's consistency. This is the "
        "standard SEEA EA output format (Table 7.1).")

    path = outdir / "03-services" / "11-filled-example.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# =====================================================================
# 4. OESA -- South Africa 2018
# =====================================================================
def generate_oesa(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Worked Example: OESA -- South Africa 2018")
    add_subtitle(doc, "Ocean Economy Satellite Account  |  Country: South Africa  |  "
                 "Year: 2018  |  Currency: ZAR (Rand) millions")
    add_note(doc, "Based on Operation Phakisa methodology (DEFF / StatsSA). Ocean partials estimated "
             "from administrative records, industry surveys, and expert judgement.")

    # ── Table 5: Scope and Partials ──
    doc.add_paragraph()
    add_title(doc, "Table 5: Scope and Ocean Partials")
    add_note(doc, "Partial = share of national industry output attributable to the ocean economy (0 to 1). "
             "RAG rating: Green = direct evidence, Amber = proxy or modelled, Red = sparse data.")

    make_table(doc,
        ["Ocean economy\ngroup", "National industry\n(examples)", "ISIC Rev. 4\ncode(s)",
         "Ocean\npartial(s)", "Source", "Confidence\n(RAG)"],
        [
            ["Living resources",
             "Fishing\nAquaculture",
             "ISIC 031\nISIC 032",
             "0.997\n0.860",
             "DEFF Fisheries\nyearbook",
             "Green"],
            ["Marine minerals",
             "Petroleum extraction\nMetal ore mining",
             "ISIC 061\nISIC 072",
             "0.989\n0.174",
             "DMRE Minerals\ndirectory",
             "Amber"],
            ["Marine construction",
             "Civil engineering\nSpecialised building",
             "ISIC 42\nISIC 43",
             "0.300\n0.378",
             "StatsSA SUT\nadjusted",
             "Amber"],
            ["Ship and boat\nbuilding",
             "Ships and boats\nmanufacture",
             "ISIC 301",
             "1.000",
             "Full allocation\n(inherently ocean)",
             "Green"],
            ["Marine transport",
             "Water transport\nSupporting transport",
             "ISIC 50\nISIC 52",
             "1.000\n0.200",
             "Transnet port\nstatistics",
             "Amber"],
            ["Coastal tourism",
             "Accommodation\nFood and beverage\nRecreation activities",
             "ISIC 55\nISIC 56\nISIC 93",
             "0.196\n0.196\n0.196",
             "SA Tourism\ncoastal survey",
             "Amber"],
        ])

    add_interpretation(doc, "Reading the partials:",
        "A partial of 0.997 for fishing means 99.7% of South Africa's fishing output is marine "
        "(the remainder is freshwater). A partial of 0.196 for coastal tourism means 19.6% of "
        "national accommodation, food, and recreation output is attributable to coastal visitors. "
        "Ship and boat building gets a full allocation (1.0) because the industry is inherently "
        "ocean-related.")

    # ── Table 6: Output, IC, GVA ──
    doc.add_page_break()
    add_title(doc, "Table 6: Ocean Economy Output, Intermediate Consumption, and GVA (R million, 2018)")

    make_table(doc,
        ["Ocean economy\ngroup", "National\noutput", "Ocean\npartial\n(weighted)", "Ocean\noutput",
         "National\nIC", "Ocean\nIC", "Ocean\nGVA", "Employment\n(persons)"],
        [
            ["Living resources", "14,252", "0.983", "14,006", "7,891", "7,757", "6,249", "27,610"],
            ["Marine minerals", "103,840", "0.684", "71,027", "48,613", "33,251", "37,776", "10,845"],
            ["Marine construction", "187,560", "0.335", "62,833", "118,563", "39,719", "23,114", "68,420"],
            ["Ship and boat\nbuilding", "4,380", "1.000", "4,380", "2,892", "2,892", "1,488", "4,216"],
            ["Marine transport", "82,940", "0.482", "39,977", "49,764", "23,986", "15,991", "31,845"],
            ["Coastal tourism", "247,300", "0.196", "48,471", "139,410", "27,324", "21,147", "124,950"],
            ["Other ocean\nindustries", "—", "—", "48,720", "—", "21,990", "26,730", "45,200"],
            ["TOTAL", "—", "—", "289,414", "—", "156,919", "132,495", "313,086"],
        ])

    add_interpretation(doc, "Interpretation:",
        "South Africa's ocean economy generated R 132,495 million in GVA in 2018. Marine minerals "
        "(mainly offshore oil and gas, and diamond mining) contributed the largest share of GVA "
        "(28.5%), followed by other ocean industries (20.2%) and living resources (4.7%). "
        "However, coastal tourism was the largest employer (124,950 persons, 39.9% of ocean "
        "employment), illustrating that GVA and employment rankings can differ substantially.")

    # ── Table 7: Summary ──
    doc.add_paragraph()
    add_title(doc, "Table 7: OESA Summary -- South Africa 2018")

    make_table(doc,
        ["Indicator", "Value", "Unit"],
        [
            ["Total ocean economy GVA", "132,495", "R million"],
            ["National GDP", "4,829,603", "R million"],
            ["Ocean economy share of GDP", "2.74", "%"],
            ["Total ocean employment", "313,086", "persons"],
            ["National employment", "16,291,000", "persons"],
            ["Ocean share of employment", "1.92", "%"],
        ])

    add_interpretation(doc, "Key finding:",
        "The ocean economy contributed 2.74% of South Africa's GDP in 2018, with 1.92% of "
        "national employment. The higher GDP share relative to employment share indicates that "
        "ocean industries (particularly minerals and transport) are more capital-intensive than "
        "the national average.")

    path = outdir / "04-oesa" / "12-filled-example.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# =====================================================================
# 5. OTSA -- Illustrative SIDS
# =====================================================================
def generate_otsa(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Worked Example: OTSA -- Illustrative Small Island Developing State")
    add_subtitle(doc, "Ocean Tourism Satellite Account  |  Country: Illustrative SIDS  |  "
                 "Year: 2023  |  Currency: USD")
    add_note(doc, "This example illustrates how to calculate coastal tourism partials when "
             "international and domestic visitors have different coastal propensities.")

    # ── Weighted partial calculation ──
    doc.add_paragraph()
    add_title(doc, "Step 1: Calculating the Weighted Expenditure Partial")

    add_interpretation(doc, "The problem:",
        "International visitors have a coastal partial of 35% (from exit surveys), while domestic "
        "tourists have a partial of 20% (from household travel surveys). A simple average (27.5%) "
        "would be wrong because international visitors spend much more per trip.")

    add_interpretation(doc, "The calculation:",
        "International visitors spend 2.5x more per trip than domestic tourists.\n\n"
        "International share of total spending:\n"
        "  = (500,000 x 2.5) / (500,000 x 2.5 + 2,000,000 x 1.0)\n"
        "  = 1,250,000 / 3,250,000 = 55.6%\n\n"
        "Domestic share of total spending:\n"
        "  = 2,000,000 / 3,250,000 = 44.4%\n\n"
        "Weighted expenditure partial:\n"
        "  = 0.556 x 0.35 + 0.444 x 0.20\n"
        "  = 0.195 + 0.089 = 0.283\n\n"
        "Rounded: 25.8% of all tourism expenditure is coastal (using spending weights)."
    )

    # ── OTSA summary table ──
    doc.add_paragraph()
    add_title(doc, "Step 2: OTSA Summary Table")

    make_table(doc,
        ["Indicator", "Total tourism\n(national)", "Coastal\npartial", "Coastal / ocean\ntourism", "Unit"],
        [
            ["International arrivals", "500,000", "0.35", "175,000", "persons/yr"],
            ["Domestic trips", "2,000,000", "0.20", "400,000", "trips/yr"],
            ["Tourism expenditure", "800", "0.258", "206", "USD million"],
            ["Tourism direct GVA", "320", "0.258", "82.5", "USD million"],
            ["Tourism employment", "45,000", "0.258", "11,610", "persons"],
            ["National GDP", "5,000", "--", "--", "USD million"],
            ["Coastal tourism % of GDP", "--", "--", "1.65", "%"],
        ])

    add_interpretation(doc, "Interpretation:",
        "Coastal and ocean tourism accounts for USD 82.5 million in GVA (1.65% of GDP) and "
        "11,610 jobs. The weighted partial (25.8%) is closer to the international partial (35%) "
        "than the domestic partial (20%) because international visitors account for a "
        "disproportionate share of spending. This method avoids double-counting by applying "
        "the partial to expenditure rather than visitor numbers.")

    add_interpretation(doc, "Sensitivity check:",
        "If the simple average partial (27.5%) had been used instead, coastal tourism GVA "
        "would have been USD 88 million -- an overestimate of 6.7%. For countries where "
        "international and domestic spending patterns are similar, the difference is smaller. "
        "For countries highly dependent on international tourism (e.g., Maldives, Seychelles), "
        "the weighted partial may be significantly different from the simple average.")

    path = outdir / "05-otsa" / "06-filled-example.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# =====================================================================
# 6. WASTE -- Coastal Country
# =====================================================================
def generate_waste(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Worked Example: Waste and Emissions -- Coastal Country")
    add_subtitle(doc, "SEEA-CF framework: tracking flows of residuals from the economy "
                 "to the marine environment  |  Year: 2022")
    add_note(doc, "Illustrative mid-income coastal country (~15 million population, "
             "mixed agricultural/industrial/service economy). Data compiled from environmental "
             "ministry reports, port authority records, and UNEP Global Wastewater Initiative.")

    # ── Table 9: Water Emissions ──
    doc.add_paragraph()
    add_title(doc, "Table 9: Water Emissions to the Marine Environment (tonnes/year)")

    make_table(doc,
        ["Source sector", "Nitrogen\n(tonnes N/yr)", "Phosphorus\n(tonnes P/yr)",
         "BOD\n(tonnes O\u2082/yr)", "Heavy metals\n(tonnes/yr)", "Total load\n(tonnes/yr)"],
        [
            ["Agriculture (fertiliser\nrunoff, livestock)", "12,400", "3,100", "8,200", "15", "23,715"],
            ["Manufacturing\n(food processing, textiles)", "1,800", "420", "6,500", "85", "8,805"],
            ["Sewerage: treated\n(secondary treatment)", "2,200", "380", "1,800", "8", "4,388"],
            ["Sewerage: untreated\n(direct discharge)", "8,600", "2,100", "18,500", "12", "29,212"],
            ["Households (direct\ncoastal discharge)", "1,400", "350", "3,200", "2", "4,952"],
            ["Aquaculture\n(fish/shrimp farms)", "2,100", "680", "4,800", "5", "7,585"],
            ["Other (stormwater,\nrecreational vessels)", "800", "120", "1,200", "3", "2,123"],
            ["TOTAL", "29,300", "7,150", "44,200", "130", "80,780"],
        ])

    add_interpretation(doc, "Interpretation:",
        "The largest pollution source by volume is untreated sewage discharge (29,212 tonnes/yr, "
        "36% of total load), highlighting the critical infrastructure gap. Agriculture is the "
        "largest source of nitrogen (12,400 tonnes N/yr, 42% of total N) due to fertiliser "
        "runoff. Manufacturing contributes the most heavy metals (85 tonnes/yr, 65%). BOD load "
        "is dominated by untreated sewage (18,500 tonnes, 42%), which drives coastal "
        "deoxygenation.")

    # ── Table 10: Solid Waste ──
    doc.add_paragraph()
    add_title(doc, "Table 10: Solid Waste Entering the Marine Environment (tonnes/year)")

    make_table(doc,
        ["Waste type", "Land-based\nsources", "Sea-based\nsources", "Total\n(tonnes/yr)", "% of total"],
        [
            ["Plastics: macroplastic\n(packaging, bags, bottles)", "8,100", "1,300", "9,400", "52%"],
            ["Plastics: microplastic\n(fibres, pellets, fragments)", "1,200", "600", "1,800", "10%"],
            ["Fishing gear\n(nets, lines, traps)", "0", "2,000", "2,000", "11%"],
            ["Ship waste\n(bilge, cargo residue)", "0", "850", "850", "5%"],
            ["Sewage sludge", "1,500", "0", "1,500", "8%"],
            ["Organic waste\n(food, vegetation)", "1,800", "200", "2,000", "11%"],
            ["Other (glass, metal,\nconstruction debris)", "350", "100", "450", "3%"],
            ["TOTAL", "12,950", "5,050", "18,000", "100%"],
        ])

    add_interpretation(doc, "Interpretation:",
        "The most persistent pollutant is plastic waste, comprising 62% of total marine solid "
        "waste (macro + micro combined: 11,200 tonnes). Land-based sources dominate (72%), "
        "but sea-based sources are significant (28%), mainly from fishing gear (2,000 tonnes, "
        "11% of total) and commercial shipping. Fishing gear, though smaller in volume than "
        "packaging plastics, causes disproportionate ecological harm through ghost fishing. "
        "The 52% macroplastic share aligns with global estimates (Jambeck et al. 2015).")

    # ── Table 11: Air Emissions ──
    doc.add_page_break()
    add_title(doc, "Table 11: Air Emissions from Ocean Industries (tonnes/year)")

    make_table(doc,
        ["Ocean sector", "Fuel consumed\n(tonnes/yr)", "CO\u2082\n(tonnes/yr)",
         "SO\u2082\n(tonnes/yr)", "NOx\n(tonnes/yr)", "PM\u2082.\u2085\n(tonnes/yr)"],
        [
            ["International shipping\n(IMO-registered)", "185,000", "573,500", "9,250", "14,800", "1,295"],
            ["Domestic shipping\nand ferries", "42,000", "130,200", "2,100", "3,360", "294"],
            ["Fishing fleet\n(motorised vessels)", "28,000", "86,800", "1,400", "2,240", "196"],
            ["Offshore oil and gas\n(platforms, supply)", "15,500", "48,050", "775", "1,240", "109"],
            ["Port operations\n(cranes, vehicles, tugs)", "8,500", "26,350", "425", "680", "60"],
            ["TOTAL", "279,000", "864,900", "13,950", "22,320", "1,954"],
        ])

    add_interpretation(doc, "Interpretation:",
        "Maritime shipping (international + domestic) dominates air emissions from the ocean "
        "economy, accounting for 81% of CO\u2082, 81% of SO\u2082, and 81% of NOx. International "
        "shipping alone contributes 66% of total ocean-sector CO\u2082. The fishing fleet, "
        "though smaller in absolute emissions, operates closer to shore and in ecologically "
        "sensitive areas. SO\u2082 emissions are a particular concern for port-adjacent communities "
        "and are subject to IMO 2020 sulphur regulations (0.50% m/m fuel sulphur limit).")

    add_interpretation(doc, "Policy linkage:",
        "These tables directly support SDG 14 (Life Below Water) indicator 14.1.1 (coastal "
        "eutrophication / floating plastic debris density) and can feed into national GHG "
        "inventory reporting under UNFCCC. The separation of treated vs. untreated sewage "
        "helps prioritise wastewater infrastructure investment.")

    path = outdir / "06-waste" / "07-filled-example.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# =====================================================================
# MAIN
# =====================================================================
def main():
    outdir = Path(__file__).parent
    print("Generating filled example accounting tables for print notebooks...")
    print()

    generate_extent(outdir)
    generate_condition(outdir)
    generate_services(outdir)
    generate_oesa(outdir)
    generate_otsa(outdir)
    generate_waste(outdir)

    print()
    print("Done! Six filled-example Word documents generated.")


if __name__ == "__main__":
    main()
