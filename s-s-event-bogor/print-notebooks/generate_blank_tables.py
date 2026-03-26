#!/usr/bin/env python3
"""Generate printable blank accounting table templates as GOAP Word docs.

One doc per account type showing the empty table structure fellows will fill in.
These go into each print-notebook alongside the filled examples.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

TEAL = RGBColor(0x0A, 0x54, 0x55)
GREEN = RGBColor(0x3B, 0x9C, 0x7B)
BODY = RGBColor(0x30, 0x30, 0x2F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_HEX = "404040"
YELLOW_HEX = "FFF2CC"


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:left w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:right w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="{GRAY_HEX}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def styled_run(p, text, size=10, bold=False, color=BODY):
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def add_title(doc, text):
    p = doc.add_paragraph()
    styled_run(p, text, size=16, color=TEAL)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    styled_run(p, text, size=11, color=GREEN)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    styled_run(p, text, size=9, color=RGBColor(0x71, 0x71, 0x71))
    p.paragraph_format.space_before = Pt(2)
    return p


def make_table(doc, headers, rows, input_cols=None):
    """Create a GOAP table with optional yellow input cells."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, h, size=10, bold=True, color=WHITE)
        set_cell_shading(cell, "3B9C7B")

    # Data rows
    for r, row_data in enumerate(rows, 1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]

            if c == 0:
                styled_run(p, str(val), size=10, bold=True, color=WHITE)
                set_cell_shading(cell, "0A5455")
            else:
                styled_run(p, str(val) if val else "", size=10, color=BODY)
                if input_cols and c in input_cols and not val:
                    set_cell_shading(cell, YELLOW_HEX)

            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    set_table_borders(table)
    return table


def set_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_w, new_h = section.page_height, section.page_width
    section.page_width = new_w
    section.page_height = new_h
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


# ===================================================================
# EXTENT
# ===================================================================
def generate_extent_tables(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Extent Account -- Blank Tables")
    add_subtitle(doc, "Fill in with your country's data")

    doc.add_paragraph()
    add_title(doc, "Table 1: Ecosystem Extent Account")
    add_note(doc, "Accounting area: ____________  |  Opening year: ______  |  Closing year: ______")
    make_table(doc,
        ["", "Ecosystem 1\n(name, code)", "Ecosystem 2\n(name, code)", "Ecosystem 3\n(name, code)", "Other", "Total"],
        [
            ["Opening extent (ha)", "", "", "", "", ""],
            ["Additions: natural expansion", "", "", "", "", ""],
            ["Additions: managed restoration", "", "", "", "", ""],
            ["Additions: reclassification", "", "", "", "", ""],
            ["Total additions", "", "", "", "", ""],
            ["Reductions: natural reduction", "", "", "", "", ""],
            ["Reductions: managed conversion", "", "", "", "", ""],
            ["Reductions: reclassification", "", "", "", "", ""],
            ["Total reductions", "", "", "", "", ""],
            ["Net change", "", "", "", "", ""],
            ["Closing extent (ha)", "", "", "", "", ""],
        ], input_cols={1, 2, 3, 4, 5})
    add_note(doc, "Total accounting area must be the same for opening and closing. Net change across all types sums to zero.")

    doc.add_page_break()
    add_title(doc, "Table 1b: Extent Time Series")
    make_table(doc,
        ["Ecosystem type", "Year 1 (ha)", "Year 2 (ha)", "Year 3 (ha)", "Year 4 (ha)", "Total change", "Change (%)"],
        [
            ["Ecosystem 1", "", "", "", "", "", ""],
            ["Ecosystem 2", "", "", "", "", "", ""],
            ["Ecosystem 3", "", "", "", "", "", ""],
            ["Other", "", "", "", "", "", ""],
            ["Total", "", "", "", "", "", ""],
        ], input_cols={1, 2, 3, 4, 5, 6})

    doc.add_paragraph()
    add_title(doc, "Table 2: Change Matrix (hectares)")
    make_table(doc,
        ["Opening \\ Closing", "Ecosystem 1", "Ecosystem 2", "Ecosystem 3", "Other", "Opening total"],
        [
            ["Ecosystem 1", "", "", "", "", ""],
            ["Ecosystem 2", "", "", "", "", ""],
            ["Ecosystem 3", "", "", "", "", ""],
            ["Other", "", "", "", "", ""],
            ["Closing total", "", "", "", "", ""],
        ], input_cols={1, 2, 3, 4, 5})
    add_note(doc, "Diagonal = stable area. Off-diagonal = transitions. Row totals = opening extent. Column totals = closing extent.")

    path = outdir / "01-extent" / "09-blank-tables.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# ===================================================================
# CONDITION
# ===================================================================
def generate_condition_tables(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Condition Account -- Blank Tables")
    add_subtitle(doc, "Fill in with your country's data")
    add_note(doc, "Ecosystem type: ____________  |  Accounting area: ____________  |  Opening year: ______  |  Closing year: ______")

    doc.add_paragraph()
    add_title(doc, "Table 3: Ecosystem Condition Account")
    make_table(doc,
        ["Indicator", "Unit", "Direction\n(higher is\nbetter/worse)", "Reference\nlevel", "Opening\nvalue",
         "Opening\nCI", "Closing\nvalue", "Closing\nCI", "Change\nin CI", "Status"],
        [
            ["Indicator 1", "", "", "", "", "", "", "", "", ""],
            ["Indicator 2", "", "", "", "", "", "", "", "", ""],
            ["Indicator 3", "", "", "", "", "", "", "", "", ""],
            ["Indicator 4", "", "", "", "", "", "", "", "", ""],
            ["Indicator 5", "", "", "", "", "", "", "", "", ""],
            ["Indicator 6", "", "", "", "", "", "", "", "", ""],
            ["Composite CI", "", "average", "", "", "", "", "", "", ""],
        ], input_cols={1, 2, 3, 4, 5, 6, 7, 8, 9})

    add_note(doc, "CI = Condition Index (0 = degraded, 1 = reference). Higher-is-better: CI = Value / Reference. Higher-is-worse: CI = 1 - (Value / Reference). Capped at 0 to 1.")

    doc.add_paragraph()
    add_title(doc, "Reference Levels")
    make_table(doc,
        ["Indicator", "Unit", "Direction", "Reference level", "Source"],
        [["", "", "", "", ""] for _ in range(6)],
        input_cols={1, 2, 3, 4})

    path = outdir / "02-condition" / "10-blank-tables.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# ===================================================================
# SERVICES
# ===================================================================
def generate_services_tables(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Ecosystem Services Account -- Blank Tables")
    add_subtitle(doc, "Part A (physical) is a complete account. Part B (monetary) is optional.")
    add_note(doc, "Accounting area: ____________  |  Year: ______")

    doc.add_paragraph()
    add_title(doc, "Table 4a: Physical Supply (Part A)")
    make_table(doc,
        ["Service", "CICES\ncategory", "Unit", "Ecosystem 1", "Ecosystem 2", "Ecosystem 3", "Total"],
        [
            ["Fish provisioning", "Provisioning", "kg/yr", "", "", "", ""],
            ["Wood provisioning", "Provisioning", "tonnes/yr", "", "", "", ""],
            ["Carbon sequestration", "Regulating", "Mg CO2/yr", "", "", "", ""],
            ["Coastal protection", "Regulating", "m coastline", "", "", "", ""],
            ["Nursery habitat", "Regulating", "kg biomass/yr", "", "", "", ""],
            ["Recreation", "Cultural", "visitors/yr", "", "", "", ""],
            ["Gleaning", "Cultural", "hours/yr", "", "", "", ""],
        ], input_cols={3, 4, 5, 6})

    doc.add_page_break()
    add_title(doc, "Table 4b: Monetary Supply (Part B) -- USD/year")
    make_table(doc,
        ["Service", "Valuation\nmethod", "Value\ntype", "Ecosystem 1", "Ecosystem 2", "Ecosystem 3", "Total"],
        [
            ["Fish provisioning", "Resource rent", "Market", "", "", "", ""],
            ["Carbon sequestration", "SCC", "Non-market", "", "", "", ""],
            ["Coastal protection", "Replacement cost", "Non-market", "", "", "", ""],
            ["Nursery habitat", "Productivity change", "Market", "", "", "", ""],
            ["Recreation", "Direct expenditure", "Market", "", "", "", ""],
            ["Gleaning", "Equivalent wage", "Mixed", "", "", "", ""],
            ["TOTAL", "", "", "", "", "", ""],
        ], input_cols={3, 4, 5, 6})
    add_note(doc, "Part A stands alone. Only fill Part B for services where you have economic data. Mixed completeness is normal.")

    path = outdir / "03-services" / "10-blank-tables.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# ===================================================================
# OESA
# ===================================================================
def generate_oesa_tables(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Ocean Economy Satellite Account -- Blank Tables")
    add_subtitle(doc, "SNA-based: reorganize national accounts to identify the ocean economy")
    add_note(doc, "Country: ____________  |  Year: ______  |  Currency: ______")

    doc.add_paragraph()
    add_title(doc, "Table 5: Scope and Ocean Partials")
    make_table(doc,
        ["Ocean economy group", "National industry", "ISIC code", "Ocean\npartial", "Source", "Confidence\n(RAG)"],
        [
            ["Living resources", "", "", "", "", ""],
            ["Marine minerals", "", "", "", "", ""],
            ["Marine construction", "", "", "", "", ""],
            ["Ship and boat building", "", "", "", "", ""],
            ["Marine transport", "", "", "", "", ""],
            ["Coastal tourism", "", "", "", "", ""],
        ], input_cols={1, 2, 3, 4, 5})
    add_note(doc, "Partial = share of industry output that is ocean-related (0 to 1). RAG: Green = direct evidence, Amber = proxy, Red = sparse.")

    doc.add_paragraph()
    add_title(doc, "Table 6: Ocean Economy Output, Intermediate Consumption, and GVA")
    make_table(doc,
        ["Ocean economy group", "National\noutput", "Ocean\npartial", "Ocean\noutput", "National\nIC",
         "Ocean\nIC", "Ocean\nGVA", "Employment"],
        [
            ["Living resources", "", "", "", "", "", "", ""],
            ["Marine minerals", "", "", "", "", "", "", ""],
            ["Marine construction", "", "", "", "", "", "", ""],
            ["Ship and boat building", "", "", "", "", "", "", ""],
            ["Marine transport", "", "", "", "", "", "", ""],
            ["Coastal tourism", "", "", "", "", "", "", ""],
            ["TOTAL", "", "", "", "", "", "", ""],
        ], input_cols={1, 2, 3, 4, 5, 6, 7})
    add_note(doc, "Ocean output = National output x Partial. Ocean IC = National IC x Partial. GVA = Output - IC.")

    doc.add_paragraph()
    add_title(doc, "Table 7: OESA Summary")
    make_table(doc,
        ["Indicator", "Value", "Unit"],
        [
            ["Total ocean economy GVA", "", ""],
            ["National GDP", "", ""],
            ["Ocean economy % of GDP", "", "%"],
            ["Total ocean employment", "", "persons"],
        ], input_cols={1})

    path = outdir / "04-oesa" / "11-blank-tables.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# ===================================================================
# OTSA
# ===================================================================
def generate_otsa_tables(outdir):
    doc = Document()

    add_title(doc, "Ocean Tourism Satellite Account -- Blank Tables")
    add_note(doc, "Country: ____________  |  Year: ______")

    doc.add_paragraph()
    add_title(doc, "Table 8: OTSA Indicators")
    make_table(doc,
        ["Indicator", "Total tourism\n(national)", "Tourism\npartial", "Coastal/ocean\ntourism", "Unit"],
        [
            ["International arrivals", "", "", "", "persons/yr"],
            ["Domestic trips", "", "", "", "trips/yr"],
            ["Tourism expenditure", "", "", "", "USD million"],
            ["Tourism direct GVA", "", "", "", "USD million"],
            ["Tourism employment", "", "", "", "persons"],
            ["National GDP", "", "", "", "USD million"],
            ["Coastal tourism % of GDP", "", "", "", "%"],
        ], input_cols={1, 2, 3})
    add_note(doc, "Coastal tourism = Total tourism x Tourism partial. Partial estimated from visitor motivation surveys, geographic proxy, or activity-based proxy.")

    path = outdir / "05-otsa" / "05-blank-tables.docx"
    doc.save(str(path))
    print(f"  {path.name}")


# ===================================================================
# WASTE
# ===================================================================
def generate_waste_tables(outdir):
    doc = Document()
    set_landscape(doc)

    add_title(doc, "Waste and Emissions Account -- Blank Tables")
    add_subtitle(doc, "SEEA-CF: tracking what the economy puts into the marine environment")
    add_note(doc, "Country: ____________  |  Year: ______")

    doc.add_paragraph()
    add_title(doc, "Table 9: Water Emissions to Sea (tonnes/year)")
    make_table(doc,
        ["Source sector", "Nitrogen\n(tonnes N/yr)", "Phosphorus\n(tonnes P/yr)", "BOD\n(tonnes/yr)",
         "Heavy metals\n(tonnes/yr)", "Total"],
        [
            ["Agriculture (runoff)", "", "", "", "", ""],
            ["Manufacturing", "", "", "", "", ""],
            ["Sewerage (treated)", "", "", "", "", ""],
            ["Sewerage (untreated)", "", "", "", "", ""],
            ["Households (direct)", "", "", "", "", ""],
            ["Other", "", "", "", "", ""],
            ["TOTAL", "", "", "", "", ""],
        ], input_cols={1, 2, 3, 4, 5})

    doc.add_paragraph()
    add_title(doc, "Table 10: Solid Waste Entering the Marine Environment (tonnes/year)")
    make_table(doc,
        ["Waste type", "Land-based\nsources", "Sea-based\nsources", "Total", "% of total"],
        [
            ["Plastics (macroplastic)", "", "", "", ""],
            ["Plastics (microplastic)", "", "", "", ""],
            ["Fishing gear", "", "", "", ""],
            ["Ship waste", "", "", "", ""],
            ["Sewage sludge", "", "", "", ""],
            ["Other", "", "", "", ""],
            ["TOTAL", "", "", "", ""],
        ], input_cols={1, 2, 3, 4})

    doc.add_paragraph()
    add_title(doc, "Table 11: Air Emissions from Ocean Industries (tonnes/year)")
    make_table(doc,
        ["Ocean sector", "Fuel\n(tonnes/yr)", "CO2\n(tonnes/yr)", "SO2\n(tonnes/yr)", "NOx\n(tonnes/yr)"],
        [
            ["International shipping", "", "", "", ""],
            ["Domestic shipping", "", "", "", ""],
            ["Fishing fleet", "", "", "", ""],
            ["Offshore oil and gas", "", "", "", ""],
            ["Port operations", "", "", "", ""],
            ["TOTAL", "", "", "", ""],
        ], input_cols={1, 2, 3, 4})

    path = outdir / "06-waste" / "06-blank-tables.docx"
    doc.save(str(path))
    print(f"  {path.name}")


def main():
    outdir = Path(__file__).parent
    print("Generating blank accounting tables for print notebooks...")
    generate_extent_tables(outdir)
    generate_condition_tables(outdir)
    generate_services_tables(outdir)
    generate_oesa_tables(outdir)
    generate_otsa_tables(outdir)
    generate_waste_tables(outdir)
    print("Done!")


if __name__ == "__main__":
    main()
