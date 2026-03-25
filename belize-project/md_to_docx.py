#!/usr/bin/env python3
"""
Convert Markdown files to GOAP-styled Word documents.

Uses the GOAP Report Template as a base and applies proper formatting
to headings, body text, tables, and captions.

Usage:
    python md_to_docx.py input.md                    # outputs input.docx
    python md_to_docx.py input.md -o report.docx     # custom output name
    python md_to_docx.py input.md -t path/to/template.docx  # custom template
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# GOAP colour palette
# ---------------------------------------------------------------------------
DARK_TEAL = RGBColor(0x0A, 0x54, 0x55)
GREEN = RGBColor(0x3B, 0x9C, 0x7B)
LIGHT_MINT = RGBColor(0xD4, 0xEE, 0xE5)
BODY_TEXT = RGBColor(0x30, 0x30, 0x2F)
GRAY_CAPTION = RGBColor(0x71, 0x71, 0x71)
BORDER_DARK_GRAY = "404040"
BORDER_GREEN = "2C745B"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

DEFAULT_TEMPLATE = (
    Path(__file__).parent
    / "Templates"
    / "GOAP Templates"
    / "GOAP Word templates"
    / "GOAP-REPORT TEMPLATE.docx"
)

# ---------------------------------------------------------------------------
# Markdown parsing helpers
# ---------------------------------------------------------------------------

def parse_markdown(text):
    """Parse markdown into a list of block elements.

    Each block is a dict with a 'type' key:
      heading  - keys: level, text
      paragraph - keys: text
      table    - keys: headers, rows
      caption  - keys: text  (lines starting with "Figure" or "Table" caption)
      blank    - empty line separator
    """
    lines = text.split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            blocks.append({
                "type": "heading",
                "level": len(m.group(1)),
                "text": m.group(2).strip(),
            })
            i += 1
            continue

        # Table (detect by pipe-delimited lines)
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(_parse_table_block(table_lines))
            continue

        # Figure/Table caption line
        stripped = line.strip()
        if stripped and (
            stripped.lower().startswith("figure")
            or stripped.lower().startswith("table ")
            or stripped.lower().startswith("source:")
            or stripped.lower().startswith("note:")
        ):
            blocks.append({"type": "caption", "text": stripped})
            i += 1
            continue

        # Blank line
        if not stripped:
            blocks.append({"type": "blank"})
            i += 1
            continue

        # Regular paragraph (collect consecutive non-blank, non-special lines)
        para_lines = []
        while i < len(lines):
            l = lines[i]
            ls = l.strip()
            if not ls:
                break
            if re.match(r"^#{1,6}\s+", l):
                break
            if "|" in l and l.strip().startswith("|"):
                break
            para_lines.append(ls)
            i += 1
        if para_lines:
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
        continue

    return blocks


def _parse_table_block(lines):
    """Parse pipe-delimited markdown table lines into headers and rows."""

    def split_row(line):
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    headers = split_row(lines[0])
    rows = []
    for line in lines[1:]:
        cells = split_row(line)
        # Skip separator rows (all dashes/colons)
        if all(re.match(r"^[-:]+$", c) for c in cells if c):
            continue
        rows.append(cells)

    return {"type": "table", "headers": headers, "rows": rows}


# ---------------------------------------------------------------------------
# Inline formatting helpers
# ---------------------------------------------------------------------------

def add_formatted_runs(paragraph, text, default_bold=False, default_color=None):
    """Add runs to a paragraph, handling **bold** and *italic* markdown."""
    # Split on bold and italic markers
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if default_color:
                run.font.color.rgb = default_color
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            if default_color:
                run.font.color.rgb = default_color
        else:
            run = paragraph.add_run(part)
            if default_bold:
                run.bold = True
            if default_color:
                run.font.color.rgb = default_color


# ---------------------------------------------------------------------------
# Table formatting
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set background fill on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(tc, **kwargs):
    """Set borders on a table cell's tcPr element.

    Keyword args are border names (top, bottom, left, right, insideH, insideV)
    mapped to dicts with keys: sz (eighths of a point), color, val (e.g. 'single').
    """
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f"<w:tcBorders {nsdecls('w')}/>")
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", 8)}" w:space="0" '
            f'w:color="{attrs.get("color", BORDER_DARK_GRAY)}"/>'
        )
        existing = tcBorders.find(qn(f"w:{edge}"))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(el)


def set_cell_margins(tc, top=0, bottom=0, left=57, right=57):
    """Set cell margins in twips. 4pt ~ 57 twips."""
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f"</w:tcMar>"
    )
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(margins)


def _is_numeric(text):
    """Check if text looks like a number."""
    try:
        float(text.replace(",", "").replace("%", "").strip())
        return True
    except ValueError:
        return False


def format_table(table, headers, rows):
    """Apply full GOAP formatting to a python-docx Table object."""
    border_def = {"sz": 8, "color": BORDER_DARK_GRAY, "val": "single"}  # 1 pt

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            # Cell margins: 4 pt top/bottom (57 twips), 4 pt left/right
            set_cell_margins(tc, top=57, bottom=57, left=57, right=57)

            # Borders on all sides
            set_cell_border(
                tc,
                top=border_def,
                bottom=border_def,
                left=border_def,
                right=border_def,
            )

            # Clear default paragraph formatting
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

            if row_idx == 0:
                # Header row
                if col_idx == 0:
                    set_cell_shading(cell, "0A5455")
                else:
                    set_cell_shading(cell, "3B9C7B")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = WHITE
            elif col_idx == 0:
                # First column in data rows (row header)
                set_cell_shading(cell, "0A5455")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = WHITE
            else:
                # Data cell
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = BODY_TEXT
                    # Right-align numeric cells
                    cell_text = cell.text.strip()
                    if _is_numeric(cell_text):
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def remove_outer_borders(table):
    """Remove heavy outer borders, keep only light internal ones."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        return
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        for edge in ["top", "left", "bottom", "right"]:
            el = borders.find(qn(f"w:{edge}"))
            if el is not None:
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_docx(blocks, template_path, output_path, header_text=None):
    """Build a .docx from parsed markdown blocks using the GOAP template."""
    doc = Document(str(template_path))

    # Update header text across all sections (default + first-page headers)
    if header_text is not None:
        for section in doc.sections:
            for hdr in (section.header, section.first_page_header, section.even_page_header):
                if hdr is None:
                    continue
                for para in hdr.paragraphs:
                    if para.text.strip():
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = header_text
                        else:
                            run = para.add_run(header_text)
                            run.font.name = "Arial"
                            run.font.size = Pt(10)
                            run.font.color.rgb = BODY_TEXT

    # Remove template placeholder content (keep styles, headers, footers)
    for para in doc.paragraphs:
        p_element = para._element
        p_element.getparent().remove(p_element)
    for table in doc.tables:
        t_element = table._tbl
        t_element.getparent().remove(t_element)

    # Map heading levels to template style names
    heading_styles = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
        5: "Heading 5",
        6: "Heading 6",
    }

    for block in blocks:
        btype = block["type"]

        if btype == "blank":
            continue

        elif btype == "heading":
            level = block["level"]
            style_name = heading_styles.get(level, "Heading 4")
            para = doc.add_paragraph(style=style_name)
            add_formatted_runs(para, block["text"])

        elif btype == "paragraph":
            para = doc.add_paragraph(style="Normal")
            add_formatted_runs(para, block["text"], default_color=BODY_TEXT)

        elif btype == "caption":
            # Use Figure style if available, else style manually
            try:
                para = doc.add_paragraph(style="Figure")
            except KeyError:
                para = doc.add_paragraph(style="Normal")
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
            run = para.add_run(block["text"])
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = GRAY_CAPTION
            run.font.name = "Arial"

        elif btype == "table":
            headers = block["headers"]
            rows = block["rows"]
            num_cols = len(headers)

            # Create table
            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            # Populate header row
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(h)
                run.font.name = "Arial"
                run.font.size = Pt(10)

            # Populate data rows
            for i, row_data in enumerate(rows):
                for j, val in enumerate(row_data):
                    if j < num_cols:
                        cell = table.rows[i + 1].cells[j]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        run = para.add_run(val)
                        run.font.name = "Arial"
                        run.font.size = Pt(10)

            # Apply GOAP formatting
            format_table(table, headers, rows)
            remove_outer_borders(table)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown to GOAP-styled Word document."
    )
    parser.add_argument("input", help="Input Markdown file")
    parser.add_argument("-o", "--output", help="Output .docx path (default: same name as input)")
    parser.add_argument(
        "--header",
        help="Header text for the document (default: derived from the first H1 title)",
    )
    parser.add_argument(
        "-t",
        "--template",
        default=str(DEFAULT_TEMPLATE),
        help="Path to GOAP Word template (default: Templates/GOAP.../GOAP-REPORT TEMPLATE.docx)",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: {input_path} not found.", file=sys.stderr)
        sys.exit(1)

    template_path = Path(args.template)
    if not template_path.exists():
        print(f"Error: template {template_path} not found.", file=sys.stderr)
        sys.exit(1)

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_suffix(".docx")

    md_text = input_path.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)

    # Derive header text: explicit flag, or first H1 title from the document
    header_text = args.header
    if header_text is None:
        for block in blocks:
            if block["type"] == "heading" and block["level"] == 1:
                header_text = block["text"]
                break

    result = build_docx(blocks, template_path, output_path, header_text=header_text)
    print(f"Created: {result}")


if __name__ == "__main__":
    main()
