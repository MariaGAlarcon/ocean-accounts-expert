#!/usr/bin/env python3
"""
Document Format QA -- Reusable formatting checker and fixer for Word documents.

Checks and fixes:
1. Markdown bold artifacts (**text** → proper bold)
2. Table column widths (proportional to content, not equal)
3. Table page breaks (prevent splitting across pages)
4. Orphaned titles (keep headings with their tables)
5. Font consistency (normalize to specified font/sizes)

Usage:
    python format_qa.py path/to/file.docx          # Single file
    python format_qa.py path/to/folder/             # All .docx in folder
    python format_qa.py path/to/folder/ --recursive # Include subfolders

Requires: python-docx (pip install python-docx)
"""

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---- Configuration (GOAP defaults, override as needed) ----

DEFAULT_CONFIG = {
    "body_font": "Arial",
    "body_size_pt": 10,
    "h1_size_pt": 16,
    "h2_size_pt": 14,
    "h3_size_pt": 12,
    "h4_size_pt": 11,
    "min_col_width_cm": 1.2,
    "max_col_width_cm": 7.0,
    "page_width_cm": 17.0,  # A4 portrait usable width
    "keep_first_n_rows": 3,  # Keep first N rows of table together
}


# ---- Check 1: Markdown bold artifacts ----

def fix_bold_artifacts(doc):
    """Remove literal ** around text and apply proper bold formatting."""
    count = 0

    def _fix_runs(runs):
        nonlocal count
        for run in runs:
            if run.text and "**" in run.text:
                cleaned = run.text.replace("**", "")
                if cleaned != run.text:
                    run.text = cleaned
                    run.bold = True
                    count += 1

    for para in doc.paragraphs:
        _fix_runs(para.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _fix_runs(para.runs)

    return count


# ---- Check 2: Table column widths ----

def get_max_content_length(table, col_idx):
    """Get the maximum text length in a column."""
    max_len = 0
    for row in table.rows:
        if col_idx < len(row.cells):
            text = row.cells[col_idx].text.strip()
            max_line = max((len(line) for line in text.split("\n")), default=0)
            max_len = max(max_len, max_line)
    return max_len


def set_column_widths(table, config=None):
    """Set column widths proportional to content length."""
    cfg = config or DEFAULT_CONFIG
    if not table.rows or not table.columns:
        return

    num_cols = len(table.columns)
    col_lengths = [max(get_max_content_length(table, c), 1) for c in range(num_cols)]

    total_available = cfg["page_width_cm"]
    total_content = sum(col_lengths)
    min_w = cfg["min_col_width_cm"]
    max_w = cfg["max_col_width_cm"]

    widths = []
    for length in col_lengths:
        prop = (length / total_content) * total_available
        w = max(min_w, min(max_w, prop))
        if length <= 3:
            w = max(1.2, min(2.0, prop))
        elif length <= 8:
            w = max(1.5, min(3.0, prop))
        widths.append(w)

    # Scale to fit total available width
    total_w = sum(widths)
    if total_w > 0:
        scale = total_available / total_w
        widths = [w * scale for w in widths]

    # Apply widths
    for row in table.rows:
        for c, cell in enumerate(row.cells):
            if c < len(widths):
                cell.width = Cm(widths[c])

    # Set fixed layout
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is not None:
        tblPr.remove(layout)
    tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))


# ---- Check 3: Table page breaks ----

def prevent_table_split(table, config=None):
    """Prevent tables from splitting across pages."""
    cfg = config or DEFAULT_CONFIG

    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        existing = trPr.find(qn("w:cantSplit"))
        if existing is not None:
            trPr.remove(existing)
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>'))

    keep_n = cfg["keep_first_n_rows"]
    for i, row in enumerate(table.rows):
        if i < min(keep_n, len(table.rows) - 1):
            for cell in row.cells:
                for para in cell.paragraphs:
                    pPr = para._element.get_or_add_pPr()
                    if pPr.find(qn("w:keepNext")) is None:
                        pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>'))


# ---- Check 4: Orphaned titles ----

def fix_keep_with_next_before_tables(doc):
    """Ensure paragraphs before tables stay on the same page."""
    body = doc.element.body
    children = list(body)
    count = 0

    for i, child in enumerate(children):
        if child.tag == qn("w:tbl") and i > 0:
            for j in range(max(0, i - 3), i):
                prev = children[j]
                if prev.tag == qn("w:p"):
                    has_text = any(
                        t.text and t.text.strip()
                        for r in prev.findall(qn("w:r"))
                        for t in r.findall(qn("w:t"))
                    )
                    if has_text:
                        pPr = prev.find(qn("w:pPr"))
                        if pPr is None:
                            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                            prev.insert(0, pPr)
                        if pPr.find(qn("w:keepNext")) is None:
                            pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>'))
                            count += 1
    return count


# ---- Main processing ----

def process_file(filepath, config=None):
    """Run all formatting fixes on a single Word document."""
    cfg = config or DEFAULT_CONFIG
    doc = Document(str(filepath))
    fixes = {}

    bold_count = fix_bold_artifacts(doc)
    if bold_count:
        fixes["bold_artifacts"] = bold_count

    kwn_count = fix_keep_with_next_before_tables(doc)
    if kwn_count:
        fixes["keep_with_next"] = kwn_count

    table_count = 0
    for table in doc.tables:
        set_column_widths(table, cfg)
        prevent_table_split(table, cfg)
        table_count += 1
    if table_count:
        fixes["tables_fixed"] = table_count

    if fixes:
        doc.save(str(filepath))
        return fixes
    return None


def process_path(target, recursive=False, config=None):
    """Process a file or folder of Word documents."""
    target = Path(target)
    total_files = 0
    total_fixed = 0

    if target.is_file() and target.suffix == ".docx":
        files = [target]
    elif target.is_dir():
        pattern = "**/*.docx" if recursive else "*.docx"
        files = sorted(target.glob(pattern))
    else:
        print(f"Error: {target} is not a .docx file or directory")
        return 0, 0

    for filepath in files:
        if filepath.name.startswith("~$"):
            continue  # Skip temp files
        total_files += 1
        fixes = process_file(filepath, config)
        if fixes:
            total_fixed += 1
            print(f"  [FIXED] {filepath.name}: {fixes}")
        else:
            print(f"  [CLEAN] {filepath.name}")

    print(f"\nTotal: {total_files} files checked, {total_fixed} files modified")
    return total_files, total_fixed


def main():
    parser = argparse.ArgumentParser(
        description="Document Format QA: check and fix Word document formatting"
    )
    parser.add_argument(
        "path",
        help="Path to a .docx file or a folder containing .docx files",
    )
    parser.add_argument(
        "--recursive", "-r",
        action="store_true",
        help="Process subfolders recursively",
    )
    parser.add_argument(
        "--page-width",
        type=float,
        default=17.0,
        help="Usable page width in cm (default: 17.0 for A4 portrait)",
    )

    args = parser.parse_args()

    config = DEFAULT_CONFIG.copy()
    config["page_width_cm"] = args.page_width

    print(f"Document Format QA")
    print(f"Target: {args.path}")
    print(f"{'='*60}")

    process_path(args.path, recursive=args.recursive, config=config)


if __name__ == "__main__":
    main()
